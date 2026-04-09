from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import get_user_model, authenticate, login as auth_login, logout
from django.utils.text import slugify
from .forms import (
    JobseekerRegisterForm, HRRegisterForm, LoginForm, ResumeUploadForm,
    EducationForm, ExperienceForm, ProjectForm, SkillForm,
    CertificateForm, ReferenceForm,
    ProfileUpdateForm, SupportTicketForm
)
from .models import JobseekerProfile, HRProfile, Resume, Education, Experience, Project, Skill, Certificate, Reference, ParsedResumeData, JobPost, ATSResult, Notification, ContactMessage
from .utils import extract_text_from_pdf, extract_text_from_docx, parse_resume_text, calculate_ats_score, calculate_general_score
from django.contrib.auth.decorators import login_required
from .decorators import hr_required, jobseeker_required, admin_required
from django.db.models import Avg, Q, Count
from django.utils import timezone

User = get_user_model()


def home(request):
    return render(request, "index.html")

def contact_view(request):
    if request.method == "POST":
        name = request.POST.get("name")
        email = request.POST.get("email")
        subject = request.POST.get("subject")
        message = request.POST.get("message")
        
        if name and email and message:
            ContactMessage.objects.create(
                name=name,
                email=email,
                subject=subject or "General Inquiry",
                message=message
            )
            messages.success(request, "Success! Your message has been sent to the CVevo team. We'll get back to you soon.")
        else:
            messages.error(request, "Please fill in all required fields.")
            
    return redirect("home")

def organizations(request):
    return render(request, "organization.html")

from django.urls import reverse

@hr_required
def hr_update_candidate_status(request):
    """Recruitment workflow: Updates status and notifies candidate."""
    if request.method == "POST":
        result_id = request.POST.get("result_id")
        action = request.POST.get("action")
        
        # Mapping button actions to DB statuses
        status_map = {
            'shortlist': 'Shortlisted',
            'interview': 'Interviewing',
            'reject': 'Rejected'
        }
        new_status = status_map.get(action)
        
        if not new_status:
            messages.error(request, "Invalid recruitment action.")
            return redirect('hr_dashboard')

        result = get_object_or_404(ATSResult.objects.select_related('resume', 'resume__jobseeker', 'job_post'), pk=result_id, job_post__hr=request.user.hr_profile)
        
        result.status = new_status
        result.save()
        
        # Notify the Jobseeker
        if result.resume.jobseeker and result.resume.jobseeker.user:
            company_name = request.user.hr_profile.company or "The Company"
            
            notif_messages = {
                'Shortlisted': f"Great news! Your application for '{result.job_post.title}' at {company_name} has been shortlisted.",
                'Interviewing': f"Interview Request: {company_name} would like to interview you for the '{result.job_post.title}' role.",
                'Rejected': f"Thank you for your interest in the '{result.job_post.title}' role at {company_name}. We have decided to move forward with other candidates."
            }
            
            notif_icons = { 'Shortlisted': 'award', 'Interviewing': 'calendar', 'Rejected': 'x-circle' }
            notif_types = { 'Shortlisted': 'success', 'Interviewing': 'info', 'Rejected': 'warning' }
            
            Notification.push(
                result.resume.jobseeker.user, 
                notif_messages[new_status], 
                icon=notif_icons[new_status], 
                notif_type=notif_types[new_status]
            )

        messages.success(request, f"Candidate status updated to {new_status}.")
        return redirect(f"{reverse('hr_candidate_detail')}?result_id={result_id}")
        
    return redirect('hr_dashboard')

@hr_required
def hr_candidate_detail(request):
    hr = request.user.hr_profile
    result_id = request.GET.get("result_id")
    
    # We fetch the ATSResult and ensure it belongs to a job post created by this HR
    result = get_object_or_404(ATSResult.objects.select_related('resume', 'job_post'), pk=result_id, job_post__hr=hr)
    
    return render(request, "pages/hr_candidate_detail.html", {
        "result": result,
    })

def resume_tips(request):
    return render(request, "resume-tips.html")

def ats_guide(request):
    return render(request, "ats-guide.html")

def page(request, template_name):
    return render(request, f"pages/{template_name}")


def register_choose(request):
    return render(request, "auth-register-choose.html")

@hr_required
def hr_resume_upload(request):
    hr = request.user.hr_profile
    job_posts = JobPost.objects.filter(hr=hr).order_by('-created_at')

    if request.method == "POST":
        job_id = request.POST.get("job_id")
        files = request.FILES.getlist("resumes")
        
        if not job_id:
            messages.error(request, "Please select a job post.")
            return redirect('hr_resume_upload')
            
        job_post = get_object_or_404(JobPost, pk=job_id, hr=hr)
        
        if job_post.status != 'Open' or job_post.is_expired:
            job_post.update_status_if_expired() # Refresh status if needed
            messages.error(request, "This job is closed and no longer accepting applications.")
            return redirect('hr_resume_upload')
        
        if not files:
            messages.error(request, "Please select at least one resume file.")
            return redirect('hr_resume_upload')

        success_count = 0
        from core.utils import calculate_ats_score, extract_text_from_pdf, extract_text_from_docx

        for f in files:
            # 1. Create Resume Object (Source: HR Bulk)
            resume = Resume.objects.create(
                jobseeker=None, # HR Bulk upload candidates are not yet users
                file=f,
                filename=f.name,
                source='HR Bulk'
            )
            
            # 2. Extract Text
            text = ""
            if f.name.endswith(".pdf"):
                text = extract_text_from_pdf(resume.file.path)
            elif f.name.endswith(".docx"):
                text = extract_text_from_docx(resume.file.path)
            
            if text:
                # 3. Analyze against the selected Job Post
                score, matched, missing, feedback = calculate_ats_score(text, job_post.requirements)
                
                # 4. Save Results
                ATSResult.objects.create(
                    resume=resume,
                    job_post=job_post,
                    score=score,
                    feedback=feedback,
                    matched_keywords=",".join(matched),
                    missing_keywords=",".join(missing)
                )
                success_count += 1
            else:
                # If extraction fails, we still have the resume object but no analysis
                pass

        messages.success(request, f"Successfully processed {success_count} resumes for '{job_post.title}'.")
        from django.urls import reverse
        return redirect(f"{reverse('hr_candidate_ranking')}?job_id={job_post.id}")

    return render(request, "pages/hr_resume_upload.html", {"job_posts": job_posts})

@hr_required
def hr_candidate_ranking(request):
    hr = request.user.hr_profile
    job_id = request.GET.get("job_id")
    job_posts = JobPost.objects.filter(hr=hr).order_by('-created_at')
    
    selected_job = None
    results = []
    
    if job_id:
        selected_job = get_object_or_404(JobPost, pk=job_id, hr=hr)
        results = ATSResult.objects.filter(job_post=selected_job).select_related('resume', 'resume__jobseeker').order_by('-score')

        # Simple Filtering
        min_score = request.GET.get("min_score")
        if min_score:
            results = results.filter(score__gte=float(min_score))

    return render(request, "pages/hr_candidate_ranking.html", {
        "job_posts": job_posts,
        "selected_job": selected_job,
        "results": results,
    })


@login_required
def profile_settings(request):
    if request.method == "POST":
        if 'update_profile' in request.POST:
            form = ProfileUpdateForm(request.POST, instance=request.user)
            if form.is_valid():
                form.save()
                messages.success(request, "Profile updated successfully.")
                return redirect('profile_settings')
        elif 'change_password' in request.POST:
            from django.contrib.auth.forms import PasswordChangeForm
            from django.contrib.auth import update_session_auth_hash
            form = PasswordChangeForm(request.user, request.POST)
            if form.is_valid():
                user = form.save()
                update_session_auth_hash(request, user)  # Important!
                messages.success(request, "Password changed successfully.")
                return redirect('profile_settings')
            else:
                for error in form.errors.values():
                    messages.error(request, error)

    return render(request, "pages/profile_settings.html")


def help_support(request):
    if request.method == "POST":
        form = SupportTicketForm(request.POST)
        if form.is_valid():
            from .models import SupportRequest
            SupportRequest.objects.create(
                user=request.user,
                topic=form.cleaned_data['topic'],
                priority=form.cleaned_data['priority'],
                subject=form.cleaned_data['subject'],
                message=form.cleaned_data['message']
            )
            messages.success(request, "Message sent successfully! Our team will get back to you soon.")
            Notification.push(request.user, "Support message sent. We'll get back to you soon.", icon="📧", notif_type="info")
            return redirect('help_support')
    return render(request, "pages/help_support.html")

@login_required
def mark_notifications_read(request):
    """Mark all notifications as read for the current user."""
    if request.method == "POST":
        request.user.notifications.filter(is_read=False).update(is_read=True)
    return redirect(request.META.get('HTTP_REFERER', '/'))

def _unique_username_from_email(email: str) -> str:
    base = slugify(email.split("@")[0]) or "user"
    username = base
    i = 1
    while User.objects.filter(username=username).exists():
        i += 1
        username = f"{base}{i}"
    return username


def register_jobseeker(request):
    is_auth = request.user.is_authenticated
    initial_data = {}
    if is_auth:
        initial_data = {
            "full_name": request.user.full_name,
            "email": request.user.email
        }

    if request.method == "POST":
        form = JobseekerRegisterForm(request.POST, user_is_authenticated=is_auth, initial=initial_data)
        if form.is_valid():
            full_name = form.cleaned_data["full_name"]
            
            if is_auth:
                user = request.user
                user.full_name = full_name
                user.role = "jobseeker"
                user.save()
                if not hasattr(user, "jobseeker_profile"):
                    JobseekerProfile.objects.create(user=user, full_name=full_name)
                    messages.success(request, "Jobseeker profile created successfully!")
                    return redirect("jobseeker_dashboard")
            else:
                email = form.cleaned_data["email"]
                user = User.objects.create_user(
                    email=email,
                    password=form.cleaned_data["password"],
                    full_name=full_name,
                    role="jobseeker"
                )
                JobseekerProfile.objects.create(user=user, full_name=full_name)
                messages.success(request, "Jobseeker account created. Please sign in.")
                return redirect("login")
    else:
        form = JobseekerRegisterForm(user_is_authenticated=is_auth, initial=initial_data)

    return render(request, "auth-register-jobseeker.html", {"form": form})


def register_hr(request):
    is_auth = request.user.is_authenticated
    initial_data = {}
    if is_auth:
        initial_data = {
            "full_name": request.user.full_name,
            "email": request.user.email
        }

    if request.method == "POST":
        form = HRRegisterForm(request.POST, user_is_authenticated=is_auth, initial=initial_data)
        if form.is_valid():
            full_name = form.cleaned_data["full_name"]
            company = form.cleaned_data["company"]
            role_title = form.cleaned_data["role"]

            if is_auth:
                # Based on rules, HR should NOT be using Google login (is_auth=True via Google)
                # But if they somehow get here, we can allow them to finish setup if they are role='hr'
                user = request.user
                user.full_name = full_name
                user.role = "hr"
                user.save()
                if not hasattr(user, "hr_profile"):
                    HRProfile.objects.create(
                        user=user,
                        full_name=full_name,
                        company=company,
                        role=role_title
                    )
                    messages.success(request, "HR profile created successfully!")
                    return redirect("hr_dashboard")
            else:
                email = form.cleaned_data["email"]
                user = User.objects.create_user(
                    email=email,
                    password=form.cleaned_data["password"],
                    full_name=full_name,
                    role="hr"
                )
                HRProfile.objects.create(
                    user=user,
                    full_name=full_name,
                    company=company,
                    role=role_title
                )
                messages.success(request, "HR account created. Please sign in.")
                return redirect("login")
    else:
        form = HRRegisterForm(user_is_authenticated=is_auth, initial=initial_data)

    return render(request, "auth-register-hr.html", {"form": form})


def login_page(request):
    if request.method == "POST":
        form = LoginForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data["email"]
            password = form.cleaned_data["password"]
            remember = form.cleaned_data["remember"]
            
            user = authenticate(request, email=email, password=password)
            if user:
                auth_login(request, user)
                if not remember:
                    request.session.set_expiry(0)
                
                if user.role == "hr":
                    return redirect("hr_dashboard")
                elif user.role == "admin" or user.is_staff:
                    return redirect("super_admin_dashboard")
                else:
                    return redirect("jobseeker_dashboard")
            else:
                messages.error(request, "Invalid email or password.")
    else:
        form = LoginForm()

    return render(request, "auth-login.html", {"form": form})


def admin_login_page(request):
    """
    Dedicated secure login page for Administrative access.
    Only users with the 'admin' role or staff privileges can log in here.
    """
    if request.user.is_authenticated:
        if request.user.role == 'admin' or request.user.is_staff:
            return redirect('super_admin_dashboard')
        else:
            # If logged in as something else, logout first to allow admin login
            logout(request)

    if request.method == "POST":
        form = LoginForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data["email"]
            password = form.cleaned_data["password"]
            remember = form.cleaned_data["remember"]
            
            user = authenticate(request, email=email, password=password)
            if user:
                if user.role == 'admin' or user.is_staff:
                    auth_login(request, user)
                    if not remember:
                        request.session.set_expiry(0)
                    messages.success(request, f"Welcome to the CVevo Command Center, {user.full_name}.")
                    return redirect("super_admin_dashboard")
                else:
                    messages.error(request, "Access denied. This portal is for administrators only.")
            else:
                messages.error(request, "Invalid administrative credentials.")
    else:
        form = LoginForm()

    return render(request, "pages/admin_login.html", {"form": form})


def logout_view(request):
    logout(request)
    messages.success(request, "Logged out successfully.")
    return redirect("login")


def google_start(request, acct):
    if acct == "hr":
        messages.error(request, "Google sign-in is not available for HR accounts.")
        return redirect("register_hr")
        
    if acct not in ("jobseeker",):
        acct = "jobseeker"
        
    request.session["oauth_account_type"] = acct
    request.session.save()
    return redirect("/accounts/google/login/")


@login_required
def post_login_redirect(request):
    u = request.user
    if hasattr(u, "hr_profile"):
        return redirect("/hr/dashboard/")
    if hasattr(u, "jobseeker_profile"):
        return redirect("/jobseeker/dashboard/")
    
    # Fallback if profile doesn't exist yet
    messages.info(request, "Please select your account type to complete your setup.")
    return redirect("/register/")


@jobseeker_required
def resume_builder(request):
    profile = request.user.jobseeker_profile
    if request.method == "POST":
        # Handle profile info (Contact / Summary)
        profile.full_name = request.POST.get("full_name", profile.full_name)
        profile.email = request.POST.get("email", profile.email)
        profile.phone = request.POST.get("phone", profile.phone)
        profile.location = request.POST.get("location", profile.location)
        profile.linkedin = request.POST.get("linkedin", profile.linkedin)
        profile.portfolio = request.POST.get("portfolio", profile.portfolio)
        profile.summary = request.POST.get("summary", profile.summary)
        profile.save()
        messages.success(request, "Profile updated.")
        Notification.push(request.user, "Contact & Summary info updated.", icon="user", notif_type="success")
        return redirect('resume_builder')

    educations = profile.educations.all()
    experiences = profile.experiences.all()
    projects = profile.projects.all()
    skills = profile.skills.all()
    certificates = profile.certificates.all()
    references = profile.references.all()

    # Split name for templates that need it
    name_parts = profile.full_name.split(' ')
    first_name = name_parts[0] if name_parts else ""
    last_name = name_parts[-1] if len(name_parts) > 1 else ""

    return render(request, "pages/resume_builder.html", {
        "educations": educations,
        "experiences": experiences,
        "projects": projects,
        "skills": skills,
        "certificates": certificates,
        "references": references,
        "first_name": first_name,
        "last_name": last_name,
    })


@jobseeker_required
def add_certificate(request):
    if request.method == "POST":
        form = CertificateForm(request.POST)
        if form.is_valid():
            cert = form.save(commit=False)
            cert.profile = request.user.jobseeker_profile
            cert.save()
            messages.success(request, "Certificate added.")
    return redirect('/resume/builder/?tab=certificates')


@jobseeker_required
def delete_certificate(request, pk):
    cert = get_object_or_404(Certificate, pk=pk, profile=request.user.jobseeker_profile)
    cert.delete()
    messages.success(request, "Certificate deleted.")
    return redirect('/resume/builder/?tab=certificates')


@jobseeker_required
def add_reference(request):
    if request.method == "POST":
        form = ReferenceForm(request.POST)
        if form.is_valid():
            ref = form.save(commit=False)
            ref.profile = request.user.jobseeker_profile
            ref.save()
            messages.success(request, "Reference added.")
    return redirect('/resume/builder/?tab=references')


@jobseeker_required
def delete_reference(request, pk):
    ref = get_object_or_404(Reference, pk=pk, profile=request.user.jobseeker_profile)
    ref.delete()
    messages.success(request, "Reference deleted.")
    return redirect('/resume/builder/?tab=references')


@jobseeker_required
def add_education(request):
    if request.method == "POST":
        form = EducationForm(request.POST)
        if form.is_valid():
            edu = form.save(commit=False)
            edu.profile = request.user.jobseeker_profile
            edu.save()
            messages.success(request, "Education added.")
    return redirect('/resume/builder/?tab=education')


@jobseeker_required
def delete_education(request, pk):
    edu = get_object_or_404(Education, pk=pk, profile=request.user.jobseeker_profile)
    edu.delete()
    messages.success(request, "Education deleted.")
    return redirect('/resume/builder/?tab=education')


@jobseeker_required
def add_experience(request):
    if request.method == "POST":
        form = ExperienceForm(request.POST)
        if form.is_valid():
            exp = form.save(commit=False)
            exp.profile = request.user.jobseeker_profile
            exp.save()
            messages.success(request, "Experience added.")
    return redirect('/resume/builder/?tab=experience')


@jobseeker_required
def delete_experience(request, pk):
    exp = get_object_or_404(Experience, pk=pk, profile=request.user.jobseeker_profile)
    exp.delete()
    messages.success(request, "Experience deleted.")
    return redirect('/resume/builder/?tab=experience')


@jobseeker_required
def add_project(request):
    if request.method == "POST":
        form = ProjectForm(request.POST)
        if form.is_valid():
            proj = form.save(commit=False)
            proj.profile = request.user.jobseeker_profile
            proj.save()
            messages.success(request, "Project added.")
    return redirect('/resume/builder/?tab=projects')


@jobseeker_required
def delete_project(request, pk):
    proj = get_object_or_404(Project, pk=pk, profile=request.user.jobseeker_profile)
    proj.delete()
    messages.success(request, "Project deleted.")
    return redirect('/resume/builder/?tab=projects')


@jobseeker_required
def add_skill(request):
    if request.method == "POST":
        form = SkillForm(request.POST)
        if form.is_valid():
            skill = form.save(commit=False)
            skill.profile = request.user.jobseeker_profile
            skill.save()
            messages.success(request, "Skill added.")
    return redirect('/resume/builder/?tab=skills')


@jobseeker_required
def delete_skill(request, pk):
    skill = get_object_or_404(Skill, pk=pk, profile=request.user.jobseeker_profile)
    skill.delete()
    messages.success(request, "Skill deleted.")
    return redirect('/resume/builder/?tab=skills')


@jobseeker_required
def analyze_resume(request, resume_id):
    resume = get_object_or_404(Resume, pk=resume_id, jobseeker=request.user.jobseeker_profile)
    # Only show 'Open' jobs that are not expired
    job_posts = JobPost.objects.filter(status='Open').order_by('-created_at')
    
    # Trigger status update for these jobs to be sure
    for job in job_posts:
        job.update_status_if_expired()
    
    # Re-fetch only truly open ones
    job_posts = job_posts.filter(status='Open')

    if request.method == "POST":
        job_id = request.POST.get("job_id")
        job_post = get_object_or_404(JobPost, pk=job_id)

        if job_post.status != 'Open' or job_post.is_expired:
            messages.error(request, "This job is closed and no longer accepting applications.")
            return redirect('analyze_resume', resume_id=resume_id)

        # Get parsed text
        parsed_data = getattr(resume, 'parsed_data', None)
        if not parsed_data:
            from .utils import extract_text_from_pdf, extract_text_from_docx
            file_path = resume.file.path
            ext = resume.filename.split('.')[-1].lower()
            text = ""
            if ext == 'pdf': text = extract_text_from_pdf(file_path)
            elif ext == 'docx': text = extract_text_from_docx(file_path)
            
            if text:
                parsed_data = ParsedResumeData.objects.create(
                    resume=resume,
                    extracted_text=text
                )
        
        if parsed_data:
            score, matched, missing, feedback = calculate_ats_score(
                parsed_data.extracted_text, 
                job_post.requirements
            )

            ATSResult.objects.create(
                resume=resume,
                job_post=job_post,
                score=score,
                feedback=feedback,
                matched_keywords=",".join(matched),
                missing_keywords=",".join(missing)
            )
            messages.success(request, f"Analysis complete for {job_post.title}. Score: {score}%")
            Notification.push(request.user, f"ATS analysis complete for '{job_post.title}': {score}% match.", icon="🎯", notif_type="success")
            return redirect('analysis_results')
        else:
            messages.error(request, "Could not extract text for analysis.")

    return render(request, "pages/analyze_resume.html", {
        "resume": resume,
        "job_posts": job_posts
    })


@jobseeker_required
def quick_analysis(request):
    resumes = Resume.objects.filter(jobseeker=request.user.jobseeker_profile).order_by('-uploaded_at')
    
    if request.method == "POST":
        resume_id = request.POST.get("resume_id")
        job_title = request.POST.get("job_title", "Quick Scan Position")
        job_description = request.POST.get("job_description")
        
        if not resume_id or not job_description:
            messages.error(request, "Please select a resume and provide a job description.")
            return redirect("quick_analysis")
            
        resume = get_object_or_404(Resume, pk=resume_id, jobseeker=request.user.jobseeker_profile)
        
        # 1. Get Resume Text
        resume_text = ""
        if resume.file.name.endswith(".pdf"):
            resume_text = extract_text_from_pdf(resume.file.path)
        elif resume.file.name.endswith(".docx"):
            resume_text = extract_text_from_docx(resume.file.path)
            
        if not resume_text:
            messages.error(request, "Could not extract text from the selected resume.")
            return redirect("quick_analysis")
            
        # 2. Run ATS Analysis — returns (score, matched, missing, feedback)
        score, matched, missing, feedback = calculate_ats_score(resume_text, job_description)
        
        # 3. Save Result
        ats_result = ATSResult.objects.create(
            resume=resume,
            job_post=None,  # Null for quick analysis
            custom_job_title=job_title,
            score=score,
            feedback=feedback,
            matched_keywords=", ".join(matched),
            missing_keywords=", ".join(missing)
        )
        
        # 4. Save score to resume for history
        resume.latest_score = score
        resume.save()
        
        messages.success(request, f"Analysis complete! Match score: {round(score)}%")
        Notification.push(request.user, f"Quick scan complete for '{job_title}': {round(score)}%.", icon="⚡", notif_type="info")
        return redirect("analysis_results")

    return render(request, "pages/quick_analysis.html", {"resumes": resumes})

@jobseeker_required
def analysis_results(request):
    results = ATSResult.objects.filter(resume__jobseeker=request.user.jobseeker_profile).order_by('-analyzed_at')
    return render(request, "pages/analysis_results.html", {"results": results})

@jobseeker_required
def delete_analysis_result(request, pk):
    result = get_object_or_404(ATSResult, pk=pk, resume__jobseeker=request.user.jobseeker_profile)
    result.delete()
    messages.success(request, "Analysis result deleted.")
    return redirect('analysis_results')


# Protected Jobseeker Views
@jobseeker_required
def jobseeker_dashboard(request):
    profile = request.user.jobseeker_profile
    resumes = profile.resumes.all().order_by('-uploaded_at')
    
    # 1. Average ATS Score
    avg_score_data = ATSResult.objects.filter(resume__jobseeker=profile).aggregate(Avg('score'))
    avg_score = round(avg_score_data['score__avg'] or 0, 1)
    
    # 2. Total Scans This Month
    now = timezone.now()
    total_scans_month = ATSResult.objects.filter(
        resume__jobseeker=profile, 
        analyzed_at__year=now.year, 
        analyzed_at__month=now.month
    ).count()
    
    # 3. Profile Strength (25% each for Edu, Exp, Skill, Proj)
    strength_val = 0
    if profile.educations.exists(): strength_val += 25
    if profile.experiences.exists(): strength_val += 25
    if profile.skills.exists(): strength_val += 25
    if profile.projects.exists(): strength_val += 25
    
    strength_label = "Low"
    if strength_val >= 100: strength_label = "Excellent"
    elif strength_val >= 75: strength_label = "Ready"
    elif strength_val >= 50: strength_label = "Moderate"
    
    # 4. Resume Scores (get latest score for each resume)
    for r in resumes:
        latest_res = ATSResult.objects.filter(resume=r).order_by('-analyzed_at').first()
        r.latest_score = latest_res.score if latest_res else None

    return render(request, "pages/jobseeker_dashboard.html", {
        "resumes": resumes,
        "avg_score": avg_score,
        "total_scans_month": total_scans_month,
        "strength_label": strength_label,
        "strength_score": strength_val,
    })


@login_required
def dashboard_search(request):
    query = request.GET.get('q', '').strip()
    results = []
    if query:
        results = JobPost.objects.filter(
            Q(title__icontains=query) | 
            Q(description__icontains=query) | 
            Q(requirements__icontains=query)
        ).order_by('-created_at')
    
    return render(request, "pages/search_results.html", {
        "query": query,
        "results": results
    })

@jobseeker_required
def resume_parse_result(request, resume_id):
    resume = get_object_or_404(Resume, pk=resume_id, jobseeker=request.user.jobseeker_profile)
    parsed_data = get_object_or_404(ParsedResumeData, resume=resume)
    
    return render(request, "pages/resume_parse_result.html", {
        "resume": resume,
        "parsed_data": parsed_data
    })


@jobseeker_required
def resume_upload(request):
    if request.method == "POST":
        files = request.FILES.getlist('resume_file')
        if not files:
            messages.error(request, "No files selected.")
            return redirect('resume_upload')
            
        success_count = 0
        last_resume_id = None
        
        for resume_file in files:
            # Check size (5MB limit)
            if resume_file.size > 5 * 1024 * 1024:
                messages.warning(request, f"Skipped '{resume_file.name}': Too large (max 5MB).")
                continue
                
            # Check extension
            ext = resume_file.name.split('.')[-1].lower()
            if ext not in ['pdf', 'docx']:
                messages.warning(request, f"Skipped '{resume_file.name}': Unsupported format.")
                continue

            resume = Resume.objects.create(
                jobseeker=request.user.jobseeker_profile,
                file=resume_file,
                filename=resume_file.name
            )
            last_resume_id = resume.id
            
            # --- Trigger Parsing ---
            file_path = resume.file.path
            text = ""
            if ext == 'pdf': text = extract_text_from_pdf(file_path)
            elif ext == 'docx': text = extract_text_from_docx(file_path)
            
            if text:
                parsed_sections = parse_resume_text(text)
                ParsedResumeData.objects.create(
                    resume=resume,
                    extracted_text=text,
                    name=parsed_sections.get("name", ""),
                    role=parsed_sections.get("role", ""),
                    email=parsed_sections.get("email", ""),
                    phone=parsed_sections.get("phone", ""),
                    skills=parsed_sections.get("skills", ""),
                    experience=parsed_sections.get("experience", ""),
                    education=parsed_sections.get("education", ""),
                )
            success_count += 1

        if success_count > 0:
            messages.success(request, f"Successfully uploaded {success_count} resume(s).")
            Notification.push(request.user, f"Uploaded {success_count} resume(s)", icon="📄", notif_type="success")
            if success_count == 1:
                return redirect('resume_parse_result', resume_id=last_resume_id)
        return redirect('resume_upload')
    else:
        form = ResumeUploadForm()
    
    resumes = request.user.jobseeker_profile.resumes.all().order_by('-uploaded_at')
    return render(request, "pages/resume_upload.html", {"form": form, "resumes": resumes})


@jobseeker_required
def general_analysis(request, resume_id):
    resume = get_object_or_404(Resume, pk=resume_id, jobseeker=request.user.jobseeker_profile)
    
    # 1. Get parsed text
    parsed_data = getattr(resume, 'parsed_data', None)
    if not parsed_data:
        # Re-parse if needed
        ext = resume.filename.split('.')[-1].lower()
        text = ""
        if ext == 'pdf': text = extract_text_from_pdf(resume.file.path)
        elif ext == 'docx': text = extract_text_from_docx(resume.file.path)
        if text:
            parsed_data = ParsedResumeData.objects.create(resume=resume, extracted_text=text)
    
    if not parsed_data or not parsed_data.extracted_text:
        messages.error(request, "Could not extract text from this resume for analysis.")
        return redirect('resume_upload')
        
    # 2. Run General Quality Scan
    scan_result = calculate_general_score(
        parsed_data.extracted_text, 
        resume.file.size, 
        resume.filename.split('.')[-1]
    )
    
    # 3. Create a result record (without a job post)
    import json
    ats_result = ATSResult.objects.create(
        resume=resume,
        job_post=None,
        custom_job_title="General Quality Scan",
        score=scan_result['quality_score'],
        feedback=json.dumps(scan_result), # Store full JSON
        matched_keywords=", ".join(scan_result['strengths']),
        missing_keywords=", ".join(scan_result['issues_found'])
    )
    
    messages.success(request, f"Quality scan complete! Structural score: {scan_result['quality_score']}%")
    return redirect('analysis_results')


@jobseeker_required
def resume_delete(request, pk):
    resume = get_object_or_404(Resume, pk=pk, jobseeker=request.user.jobseeker_profile)
    filename = resume.filename
    resume.delete()
    messages.success(request, f"Resume '{filename}' deleted.")
    return redirect('resume_upload')


@jobseeker_required
def select_template(request, template_name):
    valid_templates = [
        't1_kelly', 't2_howard', 't3_samantha_beige', 't4_samantha_white', 't5_jessie',
        't6_taylor', 't7_blue_jessie', 't8_sebastian', 't9_travis', 't10_daniel',
        't11_mira', 't12_travis_v2', 't13_kelly_v2', 't14_wes', 't15_matthew',
        't16_james', 't17_james_navy', 't18_james_sidebar', 't19_daryl', 't20_wes_v2',
        't21_daryl_clean', 't22_samantha_blue', 't23_olivia_pink'
    ]
    if template_name in valid_templates:
        profile = request.user.jobseeker_profile
        profile.selected_template = template_name
        profile.save()
        
        # Cleaner display name for messages
        display_name = template_name.split('_', 1)[-1].replace('_', ' ').title()
        messages.success(request, f"Template '{display_name}' selected successfully.")
    else:
        messages.error(request, "Invalid template selection.")
    return redirect('resume_builder')

# Protected HR Views
@hr_required
def hr_dashboard(request):
    hr = request.user.hr_profile
    
    # 0. Update status of expired jobs
    for job in JobPost.objects.filter(hr=hr, status='Open'):
        job.update_status_if_expired()

    # 1. Dashboard Stats
    open_jobs_count = JobPost.objects.filter(hr=hr, status='Open').count()
    
    # Total unique candidates who applied for this HR's jobs
    total_candidates = ATSResult.objects.filter(job_post__hr=hr).values('resume').distinct().count()
    
    # Calculate average ATS score across all results for this HR
    from django.db.models import Avg, Count, Max
    avg_score_data = ATSResult.objects.filter(job_post__hr=hr).aggregate(Avg('score'))
    avg_score = round(avg_score_data['score__avg'] or 0, 1)
    
    # Shortlisted (Score >= 80)
    shortlisted_count = ATSResult.objects.filter(job_post__hr=hr, score__gte=80).values('resume').distinct().count()

    # 2. Active Job Posts Table
    active_jobs = JobPost.objects.filter(hr=hr).annotate(
        candidate_count=Count('ats_results', distinct=True),
        top_score=Max('ats_results__score')
    ).order_by('-created_at')[:5]

    # 3. Selection Funnel Logic (Real Data)
    total_apps = ATSResult.objects.filter(job_post__hr=hr).count()
    funnel = { 'apps': 0, 'screening': 0, 'interviewing': 0 }
    
    if total_apps > 0:
        scr_count = ATSResult.objects.filter(job_post__hr=hr, status__in=['Shortlisted', 'Interviewing']).count()
        int_count = ATSResult.objects.filter(job_post__hr=hr, status='Interviewing').count()
        
        funnel = {
            'apps': 100,
            'screening': round((scr_count / total_apps) * 100),
            'interviewing': round((int_count / total_apps) * 100)
        }

    return render(request, "pages/hr_dashboard.html", {
        "open_jobs_count": open_jobs_count,
        "total_candidates": total_candidates,
        "avg_score": avg_score,
        "shortlisted_count": shortlisted_count,
        "active_jobs": active_jobs,
        "funnel": funnel,
    })

@hr_required
def hr_create_job(request):
    if request.method == "POST":
        title = request.POST.get("title")
        description = request.POST.get("description")
        required_skills = request.POST.get("required_skills")
        experience_requirements = request.POST.get("experience_requirements")
        education_requirements = request.POST.get("education_requirements")
        tools_and_technologies = request.POST.get("tools_and_technologies")
        requirements = request.POST.get("requirements")
        deadline = request.POST.get("deadline")
        
        if title and description:
            JobPost.objects.create(
                hr=request.user.hr_profile,
                title=title,
                description=description,
                required_skills=required_skills,
                experience_requirements=experience_requirements,
                education_requirements=education_requirements,
                tools_and_technologies=tools_and_technologies,
                requirements=requirements,
                deadline=deadline if deadline else None
            )
            messages.success(request, f"Job post '{title}' created successfully!")
            Notification.push(request.user, f"Job post '{title}' created.", icon="💼", notif_type="success")
            return redirect('hr_dashboard')
        else:
            messages.error(request, "Title and Description are required.")

    return render(request, "pages/hr_create_job.html")

@hr_required
def hr_manage_jobs(request):
    hr = request.user.hr_profile
    
    # Update status of expired jobs
    for job in JobPost.objects.filter(hr=hr, status='Open'):
        job.update_status_if_expired()

    from django.db.models import Count, Max
    jobs = JobPost.objects.filter(hr=hr).annotate(
        candidate_count=Count('ats_results', distinct=True),
        top_score=Max('ats_results__score')
    ).order_by('-created_at')
    
    return render(request, "pages/hr_manage_jobs.html", {"jobs": jobs})

@hr_required
def hr_delete_job(request, job_id):
    job = get_object_or_404(JobPost, pk=job_id, hr=request.user.hr_profile)
    if request.method == "POST":
        title = job.title
        job.delete()
        messages.success(request, f"Job '{title}' deleted successfully.")
    return redirect('hr_manage_jobs')


@jobseeker_required
def export_resume_docx(request):
    profile = request.user.jobseeker_profile
    from docx import Document
    from django.http import HttpResponse
    import io

    doc = Document()
    doc.add_heading(profile.full_name, 0)
    
    p = doc.add_paragraph()
    p.add_run(f"{profile.user.email} | {profile.phone}").italic = True
    if profile.location:
        p.add_run(f" | {profile.location}")

    if profile.summary:
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(profile.summary)

    doc.add_heading('Skills', level=1)
    skills = profile.skills.all()
    if skills:
        tech_skills = [f"{s.name} ({s.level})" for s in skills if s.skill_type == 'Technical']
        soft_skills = [f"{s.name} ({s.level})" for s in skills if s.skill_type == 'Soft']
        
        if tech_skills:
            doc.add_heading('Technical Skills', level=2)
            doc.add_paragraph(", ".join(tech_skills))
        if soft_skills:
            doc.add_heading('Soft Skills', level=2)
            doc.add_paragraph(", ".join(soft_skills))

    doc.add_heading('Work Experience', level=1)
    for exp in profile.experiences.all():
        doc.add_heading(f"{exp.position} at {exp.company}", level=2)
        doc.add_paragraph(f"{exp.start_date} - {exp.end_date or 'Present'}")
        doc.add_paragraph(exp.description)

    doc.add_heading('Education', level=1)
    for edu in profile.educations.all():
        doc.add_heading(f"{edu.degree} - {edu.institution}", level=2)
        doc.add_paragraph(f"{edu.start_date} - {edu.end_date or 'Present'}")

    projects = profile.projects.all()
    if projects:
        doc.add_heading('Projects', level=1)
        for proj in projects:
            doc.add_heading(proj.title, level=2)
            if proj.link:
                doc.add_paragraph(proj.link)
            doc.add_paragraph(proj.description)

    certificates = profile.certificates.all()
    if certificates:
        doc.add_heading('Training & Certificates', level=1)
        for cert in certificates:
            doc.add_heading(cert.name, level=2)
            doc.add_paragraph(f"{cert.issuer} | {cert.date_obtained or ''}")
            if cert.link: doc.add_paragraph(cert.link)

    references = profile.references.all()
    if references:
        doc.add_heading('References', level=1)
        for ref in references:
            doc.add_heading(ref.name, level=2)
            doc.add_paragraph(f"{ref.relationship} at {ref.company}")
            doc.add_paragraph(f"Email: {ref.email} | Phone: {ref.phone}")

    # Save to buffer
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)

    response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=CVevo_Resume.docx'
    
    Notification.push(request.user, "DOCX exported successfully.", icon="📥", notif_type="success")
    return response


@admin_required
def admin_edit_user(request, pk):
    """
    Dashboard-native user editing.
    """
    edit_user = get_object_or_404(User, pk=pk)
    if request.method == "POST":
        edit_user.full_name = request.POST.get("full_name", edit_user.full_name)
        edit_user.email = request.POST.get("email", edit_user.email)
        # Handle role and activity
        edit_user.role = request.POST.get("role", edit_user.role)
        edit_user.is_active = request.POST.get("is_active") == "on"
        edit_user.save()
        messages.success(request, f"User {edit_user.email} updated successfully.")
        return redirect("admin_users")
    
    return render(request, "pages/admin_user_edit.html", {"edit_user": edit_user})


@admin_required
def admin_edit_job(request, pk):
    """
    Dashboard-native job listing moderator view.
    """
    job = get_object_or_404(JobPost, pk=pk)
    applicants_count = ATSResult.objects.filter(job_post=job).count()
    
    if request.method == "POST":
        job.title = request.POST.get("title", job.title)
        job.status = request.POST.get("status", job.status)
        job.admin_note = request.POST.get("admin_note", job.admin_note)
        job.save()
        messages.success(request, f"Job listing '{job.title}' updated.")
        return redirect("admin_jobs")
    
    return render(request, "pages/admin_job_edit.html", {
        "job": job,
        "applicants_count": applicants_count
    })


@admin_required
def admin_view_resume(request, pk):
    """
    Dashboard-native resume viewing.
    """
    resume = get_object_or_404(Resume, pk=pk)
    return render(request, "pages/admin_resume_view.html", {"resume": resume})


@admin_required
def admin_create_user(request):
    """
    Dashboard-native user creation.
    """
    if request.method == "POST":
        email = request.POST.get("email")
        full_name = request.POST.get("full_name")
        role = request.POST.get("role", "jobseeker")
        password = request.POST.get("password", "CVevo@2026")
        
        if User.objects.filter(email=email).exists():
            messages.error(request, f"User with email {email} already exists.")
        else:
            User.objects.create_user(
                email=email,
                full_name=full_name,
                role=role,
                password=password,
                is_active=True
            )
            messages.success(request, f"User {email} created successfully.")
            return redirect("admin_users")
    
    return render(request, "pages/admin_user_create.html")

from .decorators import jobseeker_required, hr_required, admin_required

@admin_required
def internal_admin_dashboard(request):
    """
    CVevo Platform Command Center.
    High-level metrics for superusers to manage the system.
    """
    from django.utils import timezone
    from datetime import timedelta
    
    now = timezone.now()
    thirty_days_ago = now - timedelta(days=30)
    
    # User Stats
    total_users = User.objects.count()
    jobseekers_count = User.objects.filter(role='jobseeker').count()
    hr_count = User.objects.filter(role='hr').count()
    new_users_30d = User.objects.filter(date_joined__gte=thirty_days_ago).count()
    
    # Platform Activity
    total_resumes = Resume.objects.count()
    total_scans = ATSResult.objects.count()
    scans_30d = ATSResult.objects.filter(analyzed_at__gte=thirty_days_ago).count()
    total_jobs = JobPost.objects.count()
    
    # Calculate more metrics for Command Center
    from django.db.models import Avg, Max, Min, Count
    from collections import Counter

    avg_score = ATSResult.objects.aggregate(Avg('score'))['score__avg'] or 0
    max_score = ATSResult.objects.aggregate(Max('score'))['score__max'] or 0
    min_score = ATSResult.objects.aggregate(Min('score'))['score__min'] or 0

    # ATS Score distribution
    score_labels = ['Low (0-40)', 'Medium (40-70)', 'High (70-100)']
    low_scores = ATSResult.objects.filter(score__lt=40).count()
    med_scores = ATSResult.objects.filter(score__gte=40, score__lt=70).count()
    high_scores = ATSResult.objects.filter(score__gte=70).count()
    score_data = [low_scores, med_scores, high_scores]

    # Daily counts for charts (last 14 days for more data)
    fourteen_days_ago = now - timedelta(days=14)
    days_list = [(now - timedelta(days=i)).date() for i in range(13, -1, -1)]
    days_labels = [d.strftime("%b %d") for d in days_list]
    
    jobseeker_growth = []
    hr_growth = []
    scans_growth = []
    
    for d in days_list:
        jobseeker_growth.append(User.objects.filter(role='jobseeker', date_joined__date=d).count())
        hr_growth.append(User.objects.filter(role='hr', date_joined__date=d).count())
        scans_growth.append(ATSResult.objects.filter(analyzed_at__date=d).count())

    # Job Metrics
    # Active jobs (last 30d) vs total
    active_jobs = JobPost.objects.filter(created_at__gte=thirty_days_ago).count()
    # Jobs with no applicants
    from django.db.models import Count
    jobs_list_all = JobPost.objects.annotate(num_applicants=Count('ats_results')).order_by('-created_at')
    no_applicants_count = jobs_list_all.filter(num_applicants=0).count()
    recent_jobs_list = jobs_list_all[:50] # Show more for explorer

    # Support Stats
    from .models import SupportRequest
    pending_support_count = SupportRequest.objects.filter(is_resolved=False).count()
    all_support_list = SupportRequest.objects.select_related('user').order_by('-created_at')[:100]
    
    # Recent Platform Feed
    recent_users = User.objects.order_by('-date_joined')[:50] # Show more for explorer
    recent_scans = ATSResult.objects.select_related('resume', 'job_post').order_by('-analyzed_at')[:50]
    all_resumes = Resume.objects.select_related('jobseeker').order_by('-uploaded_at')[:100]
    
    context = {
        "stats": {
            "total_users": total_users,
            "jobseekers": jobseekers_count,
            "hr": hr_count,
            "new_users_30d": new_users_30d,
            "total_resumes": total_resumes,
            "total_scans": total_scans,
            "scans_30d": scans_30d,
            "total_jobs": total_jobs,
            "active_jobs": active_jobs,
            "no_applicants": no_applicants_count,
            "pending_support": pending_support_count,
            "avg_score": round(avg_score, 1),
            "max_score": round(max_score, 1),
            "min_score": round(min_score, 1),
        },
        "score_distribution": {
            "labels": score_labels,
            "data": score_data,
        },
        "growth_chart": {
            "labels": days_labels,
            "jobseekers": jobseeker_growth,
            "hr": hr_growth,
            "scans": scans_growth,
        },
        "all_users_list": recent_users,
        "all_jobs_list": recent_jobs_list,
        "all_resumes_list": all_resumes,
        "all_support_list": all_support_list,
        "recent_scans": recent_scans,
    }
    
    return render(request, "pages/super_admin_dashboard.html", context)

@admin_required
def admin_delete_user(request, pk):
    u = get_object_or_404(User, pk=pk)
    if u.is_superuser:
        messages.error(request, "Superusers cannot be deleted via dashboard.")
    else:
        u.delete()
        messages.success(request, f"User {u.email} deleted.")
    return redirect('super_admin_dashboard')

@admin_required
def admin_delete_job(request, pk):
    job = get_object_or_404(JobPost, pk=pk)
    job.delete()
    messages.success(request, f"Job '{job.title}' deleted.")
    return redirect('super_admin_dashboard')

@admin_required
def admin_delete_resume(request, pk):
    resume = get_object_or_404(Resume, pk=pk)
    resume.delete()
    messages.success(request, f"Resume {resume.filename} deleted.")
    return redirect('super_admin_dashboard')

@admin_required
def admin_resolve_support(request, pk):
    from .models import SupportRequest
    req = get_object_or_404(SupportRequest, pk=pk)
    req.is_resolved = True
    req.save()
    messages.success(request, "Support request marked as resolved.")
    return redirect('admin_support')

@admin_required
def admin_users_view(request):
    """
    Dedicated User Management Page with role distribution and growth stats.
    """
    users_list = User.objects.order_by('-date_joined')
    jobseekers = User.objects.filter(role='jobseeker').count()
    hr_users = User.objects.filter(role='hr').count()
    admins = User.objects.filter(role='admin').count()
    
    # Role Chart Content
    role_labels = ['Jobseekers', 'HR Users', 'Admins']
    role_counts = [jobseekers, hr_users, admins]
    
    context = {
        "users": users_list,
        "show_stats": {
            "Jobseekers": jobseekers,
            "HR": hr_users,
            "Admins": admins,
            "Total": users_list.count()
        },
        "role_chart": {
            "labels": role_labels,
            "data": role_counts
        }
    }
    return render(request, "pages/admin_users.html", context)

@admin_required
def admin_jobs_view(request):
    """
    Job Management with active/closed stats and applicant volume.
    """
    from django.db.models import Count
    jobs_list = JobPost.objects.annotate(num_applicants=Count('ats_results')).order_by('-created_at')
    active = jobs_list.filter(status='Open').count()
    closed = jobs_list.filter(status='Closed').count()
    
    # Job Status Chart
    status_labels = ['Active', 'Closed']
    status_counts = [active, closed]
    
    context = {
        "jobs": jobs_list,
        "show_stats": {
            "Active": active,
            "Closed": closed,
            "Total Listings": jobs_list.count()
        },
        "status_chart": {
            "labels": status_labels,
            "data": status_counts
        }
    }
    return render(request, "pages/admin_jobs.html", context)

@admin_required
def admin_resumes_view(request):
    """
    Resume Analytics with source breakdown and upload trends.
    """
    resumes_list = Resume.objects.select_related('jobseeker').order_by('-uploaded_at')
    
    # Source breakdown
    from django.db.models import Count
    sources = Resume.objects.values('source').annotate(count=Count('id'))
    source_labels = [s['source'] or 'Upload' for s in sources]
    source_counts = [s['count'] for s in sources]
    
    context = {
        "resumes": resumes_list,
        "show_stats": {
            "Total Resumes": resumes_list.count(),
        },
        "source_chart": {
            "labels": source_labels,
            "data": source_counts
        }
    }
    return render(request, "pages/admin_resumes.html", context)

@admin_required
def admin_support_view(request):
    """
    Support Center with priority and resolution stats.
    """
    from .models import SupportRequest
    tickets = SupportRequest.objects.select_related('user').order_by('-created_at')
    resolved = tickets.filter(is_resolved=True).count()
    pending = tickets.filter(is_resolved=False).count()
    
    context = {
        "tickets": tickets,
        "show_stats": {
            "Pending": pending,
            "Resolved": resolved,
            "Total": tickets.count()
        }
    }
    return render(request, "pages/admin_support.html", context)

@admin_required
def admin_ats_view(request):
    """
    ATS Analytics Page with score distribution and platform-wide performance.
    """
    ats_list = ATSResult.objects.select_related('resume', 'resume__jobseeker', 'job_post').order_by('-analyzed_at')
    
    # Dynamic Scoring Stats
    avg_score = ats_list.aggregate(Avg('score'))['score__avg'] or 0
    total_scans = ats_list.count()
    
    # Score Distribution (0-40, 40-70, 70-100)
    low = ats_list.filter(score__lt=40).count()
    mid = ats_list.filter(score__gte=40, score__lt=70).count()
    high = ats_list.filter(score__gte=70).count()
    
    score_labels = ['Low (0-40)', 'Mid (40-70)', 'High (70-100)']
    score_counts = [low, mid, high]
    
    context = {
        "ats_results": ats_list,
        "show_stats": {
            "Avg Score": round(avg_score, 1),
            "Total Scans": total_scans,
            "Profiles Analyzed": ats_list.values('resume__jobseeker').distinct().count()
        },
        "score_chart": {
            "labels": score_labels,
            "data": score_counts
        }
    }
    return render(request, "pages/admin_ats.html", context)


@jobseeker_required
def export_downloads(request):
    profile = request.user.jobseeker_profile
    return render(request, "pages/export_downloads.html", {"profile": profile})

from django.http import JsonResponse

@jobseeker_required
def notify_pdf_export(request):
    """AJAX endpoint to record that a user downloaded their PDF."""
    Notification.push(request.user, "PDF exported successfully", icon="file-text", notif_type="success")
    return JsonResponse({"status": "ok"})

@hr_required
def hr_reports_export(request):

    hr = request.user.hr_profile
    job_posts = JobPost.objects.filter(hr=hr).order_by('-created_at')
    
    # Optional filtering for the report preview
    job_id = request.GET.get("job_id")
    selected_job = None
    results = []
    
    if job_id:
        selected_job = get_object_or_404(JobPost, pk=job_id, hr=hr)
        results = ATSResult.objects.filter(job_post=selected_job).select_related('resume', 'resume__jobseeker').order_by('-score')

    return render(request, "pages/hr_reports_export.html", {
        "job_posts": job_posts,
        "selected_job": selected_job,
        "results": results,
    })

@hr_required
def hr_export_csv(request):
    import csv
    from django.http import HttpResponse
    
    hr = request.user.hr_profile
    job_id = request.GET.get("job_id")
    
    if not job_id:
        messages.error(request, "Please select a job post to export.")
        return redirect('hr_reports_export')
        
    job_post = get_object_or_404(JobPost, pk=job_id, hr=hr)
    results = ATSResult.objects.filter(job_post=job_post).select_related('resume', 'resume__jobseeker').order_by('-score')
    
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{job_post.title}_Rankings.csv"'
    
    writer = csv.writer(response)
    writer.writerow(['Rank', 'Candidate Name', 'ATS Score', 'Matched Skills', 'Missing Skills', 'Analyzed At'])
    
    for i, res in enumerate(results, 1):
        name = res.resume.jobseeker.full_name if res.resume.jobseeker else f"Candidate {res.resume.id}"
        writer.writerow([
            i,
            name,
            res.score,
            res.matched_keywords,
            res.missing_keywords,
            res.analyzed_at.strftime("%Y-%m-%d %H:%M")
        ])
        
    return response


@admin_required
def admin_delete_support(request, pk):
    """Delete a support ticket."""
    ticket = get_object_or_404(SupportRequest, pk=pk)
    ticket.delete()
    messages.success(request, "Support ticket deleted.")
    return redirect("admin_support")