from django.shortcuts import get_object_or_404
from django.http import FileResponse
import json
from io import BytesIO
from datetime import datetime, timedelta
from django.contrib.auth.password_validation import validate_password
from django.core.exceptions import ValidationError
from django.core.validators import validate_email as django_validate_email
from django.utils.decorators import method_decorator
from django.views.decorators.csrf import csrf_exempt

from rest_framework import viewsets, permissions, status, authentication
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.authentication import SessionAuthentication
from rest_framework.decorators import action

class CsrfExemptSessionAuthentication(SessionAuthentication):
    def enforce_csrf(self, request):
        return  # To not perform the csrf check previously happening

from .utils import (
    extract_text_from_pdf, extract_text_from_docx,
    parse_resume_text, calculate_ats_score, calculate_general_score
)
from .models import (
    User, JobseekerProfile, HRProfile, JobPost, 
    Resume, ParsedResumeData, ATSResult, 
    Notification, SupportRequest, ContactMessage,
    Education, Experience, Project, Skill, Certificate, Reference
)
from .serializers import (
    UserSerializer, JobseekerProfileSerializer, HRProfileSerializer, 
    JobPostSerializer, ResumeSerializer, ATSResultSerializer, 
    NotificationSerializer, SupportRequestSerializer, ContactMessageSerializer
)
from .utils import extract_text_from_pdf, extract_text_from_docx, parse_resume_text

import json
import docx

NOTIF_SOUND_PRIORITY = {"high"}
NOTIF_VISIBLE_LIMIT = 5

def _notify(user, *, title, message, icon="info", notif_type="info", priority="medium", category="general", target_role=None, action_url="", metadata=None):
    return Notification.push(
        user=user,
        title=title,
        message=message,
        icon=icon,
        notif_type=notif_type,
        priority=priority,
        category=category,
        target_role=target_role,
        action_url=action_url,
        metadata=metadata or {},
    )

def _notify_jobseeker(user, **kwargs):
    return _notify(user, target_role="jobseeker", **kwargs)

def _notify_hr(user, **kwargs):
    return _notify(user, target_role="hr", **kwargs)


def _required_field_response(field_name):
    pretty_name = field_name.replace('_', ' ').strip().capitalize()
    return Response({ 'error': f'{pretty_name} is required.' }, status=400)


def _valid_email_response(email):
    if not email:
        return _required_field_response('email')
    try:
        django_validate_email(email)
    except ValidationError:
        return Response({'error': 'Please enter a valid email address.'}, status=400)
    return None


def _weak_password_response(password, user=None):
    if not password:
        return _required_field_response('password')
    try:
        validate_password(password, user=user)
    except ValidationError as exc:
        message = ' '.join(exc.messages) if getattr(exc, 'messages', None) else 'Password is weak. Please use another password.'
        return Response(
            {
                'error': message,
                'detail': message,
                'message': message,
            },
            status=400
        )
    return None


def _job_to_jd_fields(job):
    if not job:
        return None
    return {
        "title": job.title or "",
        "description": job.description or "",
        "required_skills": job.required_skills or "",
        "experience_requirements": job.experience_requirements or "",
        "education_requirements": job.education_requirements or "",
        "tools_and_technologies": job.tools_and_technologies or "",
        "requirements": job.requirements or "",
    }

# ==========================
# AUTH & USER VIEWS
# ==========================
@method_decorator(csrf_exempt, name='dispatch')
class AuthView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.AllowAny]
    def post(self, request, *args, **kwargs):
        try:
            action_type = kwargs.get('action') or request.data.get('action') or request.query_params.get('action')
            
            if action_type == 'login':
                from django.contrib.auth import authenticate, login
                email = request.data.get('email')
                password = request.data.get('password')
                user = authenticate(request, username=email, password=password)
                if user:
                    login(request, user)
                    return Response({
                        'status': 'success', 
                        'user': {
                            'id': user.id,
                            'email': user.email,
                            'full_name': user.full_name,
                            'role': user.role
                        }
                    })
                return Response({'error': 'Invalid credentials'}, status=401)
            
            elif action_type == 'register-jobseeker':
                email = (request.data.get('email') or '').strip()
                password = request.data.get('password') or ''
                full_name = (request.data.get('full_name') or '').strip()
                if not full_name:
                    return _required_field_response('full_name')
                invalid_email = _valid_email_response(email)
                if invalid_email:
                    return invalid_email
                if not password:
                    return _required_field_response('password')
                if User.objects.filter(email=email).exists():
                    return Response({'error': 'Email already exists. Please use another email.'}, status=400)
                weak_password = _weak_password_response(
                    password,
                    user=User(email=email, full_name=full_name, role='jobseeker')
                )
                if weak_password:
                    return weak_password
                user = User.objects.create_user(email=email, password=password, full_name=full_name, role='jobseeker')
                JobseekerProfile.objects.get_or_create(
                    user=user,
                    defaults={'full_name': full_name, 'email': email}
                )
                from django.contrib.auth import login
                login(request, user, backend='django.contrib.auth.backends.ModelBackend')
                return Response({
                    'status': 'success', 
                    'user': {
                        'id': user.id,
                        'email': user.email,
                        'full_name': user.full_name,
                        'role': 'jobseeker'
                    }
                })

            elif action_type == 'register-hr':
                email = (request.data.get('email') or '').strip()
                password = request.data.get('password') or ''
                full_name = (request.data.get('full_name') or '').strip()
                if not full_name:
                    return _required_field_response('full_name')
                invalid_email = _valid_email_response(email)
                if invalid_email:
                    return invalid_email
                if not password:
                    return _required_field_response('password')
                company = request.data.get('company', 'Company')
                if User.objects.filter(email=email).exists():
                    return Response({'error': 'Email already exists. Please use another email.'}, status=400)
                role_title = request.data.get('role_title') or request.data.get('role') or 'HR Manager'
                weak_password = _weak_password_response(
                    password,
                    user=User(email=email, full_name=full_name, role='hr')
                )
                if weak_password:
                    return weak_password
                user = User.objects.create_user(email=email, password=password, full_name=full_name, role='hr')
                HRProfile.objects.get_or_create(user=user, defaults={'full_name': full_name, 'company': company, 'role': role_title})
                from django.contrib.auth import login
                login(request, user, backend='django.contrib.auth.backends.ModelBackend')
                return Response({
                    'status': 'success', 
                    'user': {
                        'id': user.id,
                        'email': user.email,
                        'full_name': user.full_name,
                        'role': 'hr'
                    }
                })

            elif action_type == 'logout':
                from django.contrib.auth import logout
                logout(request)
                return Response({'status': 'success'})

            return Response({'error': 'Invalid action'}, status=400)
        except Exception as e:
            return Response({'error': str(e)}, status=500)


@method_decorator(csrf_exempt, name='dispatch')
class AdminLoginView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.AllowAny]

    def post(self, request):
        from django.contrib.auth import authenticate, login

        email = request.data.get('email')
        password = request.data.get('password')

        user = authenticate(request, username=email, password=password)
        if not user:
            return Response({'error': 'Invalid credentials'}, status=401)

        if user.role != 'admin' and not user.is_staff and not user.is_superuser:
            return Response({'error': 'Admin access only.'}, status=403)

        login(request, user)
        return Response({
            'status': 'success',
            'user': {
                'id': user.id,
                'email': user.email,
                'full_name': user.full_name,
                'role': 'admin'
            }
        })


@method_decorator(csrf_exempt, name='dispatch')
class AdminDashboardView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def _is_admin(self, user):
        return bool(user and user.is_authenticated and (user.role == 'admin' or user.is_staff or user.is_superuser))

    def get(self, request):
        from django.db.models import Avg, Max, Min, Count
        from django.utils import timezone
        from django.utils.timesince import timesince

        if not self._is_admin(request.user):
            return Response({'error': 'Admin access only.'}, status=403)

        now = timezone.localtime(timezone.now())
        total_users = User.objects.count()
        jobseekers = User.objects.filter(role='jobseeker').count()
        hr_users = User.objects.filter(role='hr').count()
        admin_users = User.objects.filter(role='admin').count()
        total_resumes = Resume.objects.count()
        total_scans = ATSResult.objects.count()
        active_jobs = JobPost.objects.filter(status='Open').count()
        total_jobs = JobPost.objects.count()
        new_users_30d = User.objects.filter(date_joined__gte=now - timedelta(days=30)).count()
        scans_30d = ATSResult.objects.filter(analyzed_at__gte=now - timedelta(days=30)).count()
        pending_support = SupportRequest.objects.filter(is_resolved=False).count() + ContactMessage.objects.filter(is_resolved=False).count()

        avg_score_data = ATSResult.objects.aggregate(avg=Avg('score'), max_score=Max('score'), min_score=Min('score'))
        avg_score = round(avg_score_data['avg'] or 0, 1)
        max_score = round(avg_score_data['max_score'] or 0, 1)
        min_score = round(avg_score_data['min_score'] or 0, 1)

        score_buckets = {
            'Below 50': ATSResult.objects.filter(score__lt=50).count(),
            '50 - 79': ATSResult.objects.filter(score__gte=50, score__lt=80).count(),
            '80+': ATSResult.objects.filter(score__gte=80).count(),
        }

        growth_labels = []
        growth_jobseekers = []
        growth_hr = []
        growth_scans = []
        for offset in range(5, -1, -1):
            total_months = now.year * 12 + now.month - 1 - offset
            year = total_months // 12
            month = total_months % 12 + 1
            label = datetime(year, month, 1).strftime('%b %Y')
            growth_labels.append(label)
            growth_jobseekers.append(User.objects.filter(role='jobseeker', date_joined__year=year, date_joined__month=month).count())
            growth_hr.append(User.objects.filter(role='hr', date_joined__year=year, date_joined__month=month).count())
            growth_scans.append(ATSResult.objects.filter(analyzed_at__year=year, analyzed_at__month=month).count())

        recent_scans = []
        for result in ATSResult.objects.select_related('resume', 'job_post', 'resume__jobseeker').order_by('-analyzed_at')[:6]:
            job_title = result.job_post.title if result.job_post else result.custom_job_title or 'Quick Scan'
            owner = result.resume.jobseeker.full_name if result.resume and result.resume.jobseeker else 'Candidate'
            recent_scans.append({
                'filename': result.resume.filename if result.resume else 'Resume',
                'job_title': job_title,
                'owner': owner,
                'score': round(result.score or 0, 1),
                'time_ago': f"{timesince(result.analyzed_at, now)} ago" if result.analyzed_at else '',
            })

        jobs = []
        for job in JobPost.objects.select_related('hr').order_by('-created_at')[:5]:
            jobs.append({
                'title': job.title,
                'company': job.hr.company if job.hr else '',
                'num_applicants': ATSResult.objects.filter(job_post=job).count(),
            })

        return Response({
            'stats': {
                'total_users': total_users,
                'new_users_30d': new_users_30d,
                'jobseekers': jobseekers,
                'hr': hr_users,
                'admins': admin_users,
                'total_resumes': total_resumes,
                'total_scans': total_scans,
                'scans_30d': scans_30d,
                'active_jobs': active_jobs,
                'total_jobs': total_jobs,
                'pending_support': pending_support,
                'avg_score': avg_score,
                'max_score': max_score,
                'min_score': min_score,
            },
            'growth_chart': {
                'labels': growth_labels,
                'jobseekers': growth_jobseekers,
                'hr': growth_hr,
                'scans': growth_scans,
            },
            'score_distribution': {
                'labels': list(score_buckets.keys()),
                'data': list(score_buckets.values()),
            },
            'recent_scans': recent_scans,
            'jobs': jobs,
        })


@method_decorator(csrf_exempt, name='dispatch')
class AdminUsersView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def _is_admin(self, user):
        return bool(user and user.is_authenticated and (user.role == 'admin' or user.is_staff or user.is_superuser))

    def _serialize_user(self, user):
        company = ''
        if hasattr(user, 'hr_profile'):
            company = user.hr_profile.company or ''
        return {
            'id': user.id,
            'full_name': user.full_name,
            'email': user.email,
            'role': user.role,
            'date_joined': user.date_joined.strftime('%b %d, %Y') if getattr(user, 'date_joined', None) else '',
            'is_active': user.is_active,
            'is_verified': user.is_verified,
            'company': company,
        }

    def get(self, request, pk=None):
        if not self._is_admin(request.user):
            return Response({'error': 'Admin access only.'}, status=403)

        if pk is not None:
            user = get_object_or_404(User, id=pk)
            return Response(self._serialize_user(user))

        users = User.objects.all().order_by('-date_joined')
        role_labels = ['jobseeker', 'hr', 'admin']
        role_counts = {role: User.objects.filter(role=role).count() for role in role_labels}

        return Response({
            'users': [self._serialize_user(user) for user in users[:200]],
            'show_stats': [
                {'label': 'Total Users', 'value': User.objects.count()},
                {'label': 'Jobseekers', 'value': role_counts['jobseeker']},
                {'label': 'HR Users', 'value': role_counts['hr']},
                {'label': 'Admins', 'value': role_counts['admin']},
            ],
            'role_chart': {
                'labels': ['Jobseekers', 'HR', 'Admins'],
                'data': [role_counts['jobseeker'], role_counts['hr'], role_counts['admin']],
            }
        })

    def post(self, request):
        if not self._is_admin(request.user):
            return Response({'error': 'Admin access only.'}, status=403)

        email = (request.data.get('email') or '').strip()
        password = request.data.get('password') or ''
        full_name = (request.data.get('full_name') or '').strip()
        role = (request.data.get('role') or 'jobseeker').strip()

        if not email or not password or not full_name:
            return Response({'error': 'Full name, email, and password are required.'}, status=400)
        invalid_email = _valid_email_response(email)
        if invalid_email:
            return invalid_email
        if User.objects.filter(email=email).exists():
            return Response({'error': 'Email already exists. Please use another email.'}, status=400)
        weak_password = _weak_password_response(
            password,
            user=User(email=email, full_name=full_name, role=role)
        )
        if weak_password:
            return weak_password

        user = User.objects.create_user(email=email, password=password, full_name=full_name, role=role)

        if role == 'admin':
            user.is_staff = True
            user.is_superuser = True
            user.save(update_fields=['is_staff', 'is_superuser'])
        elif role == 'hr':
            company = request.data.get('company') or full_name
            HRProfile.objects.get_or_create(
                user=user,
                defaults={'full_name': full_name, 'company': company, 'role': request.data.get('role_title') or 'HR Manager'}
            )
        else:
            JobseekerProfile.objects.get_or_create(user=user, defaults={'full_name': full_name, 'email': email})

        return Response(self._serialize_user(user), status=201)

    def put(self, request, pk):
        if not self._is_admin(request.user):
            return Response({'error': 'Admin access only.'}, status=403)

        user = get_object_or_404(User, id=pk)
        full_name = (request.data.get('full_name') or '').strip()
        email = (request.data.get('email') or '').strip()
        role = (request.data.get('role') or user.role).strip()
        is_active = request.data.get('is_active')
        is_verified = request.data.get('is_verified')

        if email and User.objects.exclude(id=user.id).filter(email=email).exists():
            return Response({'error': 'Email already exists.'}, status=400)

        if full_name:
            user.full_name = full_name
        if email:
            user.email = email
        user.role = role
        if is_active is not None:
            user.is_active = bool(is_active)
        if is_verified is not None:
            user.is_verified = bool(is_verified)

        if role == 'admin':
            user.is_staff = True
            user.is_superuser = True
        else:
            user.is_staff = False
            user.is_superuser = False

        user.save()

        if role == 'hr' and hasattr(user, 'hr_profile'):
            profile = user.hr_profile
            profile.full_name = full_name or profile.full_name
            profile.company = request.data.get('company') or profile.company or profile.full_name
            profile.role = request.data.get('role_title') or profile.role
            profile.save()
        elif role == 'jobseeker' and hasattr(user, 'jobseeker_profile'):
            profile = user.jobseeker_profile
            profile.full_name = full_name or profile.full_name
            profile.email = email or profile.email
            profile.save()

        return Response(self._serialize_user(user))

    def delete(self, request, pk):
        if not self._is_admin(request.user):
            return Response({'error': 'Admin access only.'}, status=403)
        if request.user.id == pk:
            return Response({'error': 'You cannot delete your own account.'}, status=400)

        user = get_object_or_404(User, id=pk)
        user.delete()
        return Response({'status': 'deleted'})


@method_decorator(csrf_exempt, name='dispatch')
class AdminJobsView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def _is_admin(self, user):
        return bool(user and user.is_authenticated and (user.role == 'admin' or user.is_staff or user.is_superuser))

    def _serialize_job(self, job):
        from django.db.models import Max
        applicants_count = ATSResult.objects.filter(job_post=job).count()
        return {
            'id': job.id,
            'title': job.title,
            'company': job.hr.company if job.hr else '',
            'hr_name': job.hr.full_name if job.hr else '',
            'created_at': job.created_at.strftime('%b %d, %Y') if job.created_at else '',
            'status': job.status,
            'applicants_count': applicants_count,
            'num_applicants': applicants_count,
            'admin_note': job.admin_note or '',
            'top_score': ATSResult.objects.filter(job_post=job).aggregate(max_score=Max('score'))['max_score'] or 0,
        }

    def get(self, request, pk=None):
        from django.db.models import Count
        if not self._is_admin(request.user):
            return Response({'error': 'Admin access only.'}, status=403)

        if pk is not None:
            job = get_object_or_404(JobPost, id=pk)
            data = self._serialize_job(job)
            data['description'] = job.description
            data['admin_note'] = job.admin_note or ''
            return Response(data)

        jobs = JobPost.objects.select_related('hr').order_by('-created_at')
        status_counts = {
            'Open': jobs.filter(status='Open').count(),
            'Closed': jobs.filter(status='Closed').count(),
            'Disabled': jobs.filter(status='Disabled').count(),
        }

        return Response({
            'jobs': [self._serialize_job(job) for job in jobs[:200]],
            'show_stats': [
                {'label': 'Total Jobs', 'value': jobs.count()},
                {'label': 'Open Jobs', 'value': status_counts['Open']},
                {'label': 'Closed Jobs', 'value': status_counts['Closed']},
                {'label': 'Disabled Jobs', 'value': status_counts['Disabled']},
            ],
            'status_chart': {
                'labels': ['Open', 'Closed', 'Disabled'],
                'data': [status_counts['Open'], status_counts['Closed'], status_counts['Disabled']],
            }
        })

    def put(self, request, pk):
        if not self._is_admin(request.user):
            return Response({'error': 'Admin access only.'}, status=403)
        job = get_object_or_404(JobPost, id=pk)

        title = (request.data.get('title') or '').strip()
        status_value = (request.data.get('status') or job.status).strip()
        admin_note = request.data.get('admin_note')

        if title:
            job.title = title
        if status_value in {'Open', 'Closed', 'Disabled'}:
            job.status = status_value
        if admin_note is not None:
            job.admin_note = admin_note

        job.save()
        return Response(self._serialize_job(job))

    def delete(self, request, pk):
        if not self._is_admin(request.user):
            return Response({'error': 'Admin access only.'}, status=403)
        job = get_object_or_404(JobPost, id=pk)
        job.delete()
        return Response({'status': 'deleted'})


@method_decorator(csrf_exempt, name='dispatch')
class AdminResumesView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def _is_admin(self, user):
        return bool(user and user.is_authenticated and (user.role == 'admin' or user.is_staff or user.is_superuser))

    def _serialize_resume(self, resume):
        owner = resume.jobseeker.full_name if resume.jobseeker else 'Bulk Upload'
        return {
            'id': resume.id,
            'filename': resume.filename,
            'owner': owner,
            'source': resume.source,
            'uploaded_at': resume.uploaded_at.strftime('%b %d, %Y') if resume.uploaded_at else '',
            'file_url': resume.file.url if resume.file else '',
        }

    def get(self, request, pk=None):
        from django.db.models import Avg
        if not self._is_admin(request.user):
            return Response({'error': 'Admin access only.'}, status=403)

        if pk is not None:
            resume = get_object_or_404(Resume, id=pk)
            return Response(self._serialize_resume(resume))

        resumes = Resume.objects.select_related('jobseeker').order_by('-uploaded_at')
        source_counts = {
            'Jobseeker': resumes.filter(source='Jobseeker').count(),
            'HR Bulk': resumes.filter(source='HR Bulk').count(),
        }
        return Response({
            'resumes': [self._serialize_resume(resume) for resume in resumes[:200]],
            'show_stats': [
                {'label': 'Total Resumes', 'value': resumes.count()},
                {'label': 'Jobseeker Uploads', 'value': source_counts['Jobseeker']},
                {'label': 'Bulk Uploads', 'value': source_counts['HR Bulk']},
            ],
            'source_chart': {
                'labels': ['Jobseeker', 'HR Bulk'],
                'data': [source_counts['Jobseeker'], source_counts['HR Bulk']],
            }
        })

    def delete(self, request, pk):
        if not self._is_admin(request.user):
            return Response({'error': 'Admin access only.'}, status=403)
        resume = get_object_or_404(Resume, id=pk)
        resume.delete()
        return Response({'status': 'deleted'})


@method_decorator(csrf_exempt, name='dispatch')
class AdminATSView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def _is_admin(self, user):
        return bool(user and user.is_authenticated and (user.role == 'admin' or user.is_staff or user.is_superuser))

    def _serialize_result(self, result):
        skill_count = len(result.matched_list)
        job_title = result.job_post.title if result.job_post else result.custom_job_title or ''
        jobseeker = result.resume.jobseeker.full_name if result.resume and result.resume.jobseeker else 'Bulk Upload'
        return {
            'id': result.id,
            'jobseeker': jobseeker,
            'target_job': job_title,
            'score': round(result.score or 0, 1),
            'skills_found': skill_count,
            'scanned_at': result.analyzed_at.strftime('%b %d, %Y') if result.analyzed_at else '',
        }

    def get(self, request, pk=None):
        from django.db.models import Avg

        if not self._is_admin(request.user):
            return Response({'error': 'Admin access only.'}, status=403)

        if pk is not None:
            result = get_object_or_404(ATSResult, id=pk)
            return Response(self._serialize_result(result))

        results = ATSResult.objects.select_related('resume', 'job_post', 'resume__jobseeker').order_by('-analyzed_at')
        score_buckets = {
            'Below 50': results.filter(score__lt=50).count(),
            '50 - 79': results.filter(score__gte=50, score__lt=80).count(),
            '80+': results.filter(score__gte=80).count(),
        }
        return Response({
            'ats_results': [self._serialize_result(result) for result in results[:200]],
            'show_stats': [
                {'label': 'Total Scans', 'value': results.count()},
                {'label': 'Avg Score', 'value': f"{round(results.aggregate(avg=Avg('score'))['avg'] or 0, 1)}%"},
                {'label': 'High Scores', 'value': score_buckets['80+']},
            ],
            'score_chart': {
                'labels': list(score_buckets.keys()),
                'data': list(score_buckets.values()),
            }
        })

    def delete(self, request, pk):
        if not self._is_admin(request.user):
            return Response({'error': 'Admin access only.'}, status=403)
        result = get_object_or_404(ATSResult, id=pk)
        result.delete()
        return Response({'status': 'deleted'})


@method_decorator(csrf_exempt, name='dispatch')
class AdminSupportView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def _is_admin(self, user):
        return bool(user and user.is_authenticated and (user.role == 'admin' or user.is_staff or user.is_superuser))

    def _format_dt(self, value):
        return value.strftime('%b %d, %Y') if value else ''

    def _serialize_support_request(self, ticket):
        requester_name = ''
        if ticket.user:
            requester_name = ticket.user.full_name or ticket.user.email
        return {
            'id': ticket.id,
            'kind': 'support',
            'source_label': 'Support Request',
            'subject': ticket.subject,
            'message': ticket.message,
            'requester_name': requester_name,
            'requester_email': ticket.user.email if ticket.user else '',
            'user_name': requester_name,
            'is_resolved': ticket.is_resolved,
            'created_at': self._format_dt(ticket.created_at),
            'priority': ticket.priority,
        }

    def _serialize_contact_message(self, message):
        return {
            'id': message.id,
            'kind': 'contact',
            'source_label': 'Contact Message',
            'subject': message.subject,
            'message': message.message,
            'requester_name': message.name or 'Website Visitor',
            'requester_email': message.email or '',
            'user_name': message.name or 'Website Visitor',
            'is_resolved': message.is_resolved,
            'created_at': self._format_dt(message.created_at),
            'priority': 'Normal',
        }

    def _get_ticket(self, pk, kind=None):
        if kind == 'contact':
            return get_object_or_404(ContactMessage, id=pk), 'contact'
        if kind == 'support':
            return get_object_or_404(SupportRequest, id=pk), 'support'

        try:
            return get_object_or_404(SupportRequest, id=pk), 'support'
        except Exception:
            return get_object_or_404(ContactMessage, id=pk), 'contact'

    def get(self, request, pk=None):
        if not self._is_admin(request.user):
            return Response({'error': 'Admin access only.'}, status=403)

        kind = request.query_params.get('kind')
        if pk is not None:
            ticket, ticket_kind = self._get_ticket(pk, kind)
            if ticket_kind == 'support':
                return Response(self._serialize_support_request(ticket))
            return Response(self._serialize_contact_message(ticket))

        support_requests = list(SupportRequest.objects.select_related('user').order_by('-created_at'))
        contact_messages = list(ContactMessage.objects.order_by('-created_at'))
        combined = (
            [('support', ticket.created_at, ticket) for ticket in support_requests] +
            [('contact', message.created_at, message) for message in contact_messages]
        )
        combined.sort(key=lambda item: item[1] or datetime.min, reverse=True)
        tickets = [
            self._serialize_support_request(item[2]) if item[0] == 'support'
            else self._serialize_contact_message(item[2])
            for item in combined
        ]

        return Response({
            'tickets': tickets[:200],
            'show_stats': [
                {
                    'label': 'Open Tickets',
                    'value': SupportRequest.objects.filter(is_resolved=False).count() + ContactMessage.objects.filter(is_resolved=False).count()
                },
                {
                    'label': 'Resolved',
                    'value': SupportRequest.objects.filter(is_resolved=True).count() + ContactMessage.objects.filter(is_resolved=True).count()
                },
                {'label': 'Urgent', 'value': SupportRequest.objects.filter(priority='Urgent').count()},
            ],
        })

    def post(self, request, pk):
        if not self._is_admin(request.user):
            return Response({'error': 'Admin access only.'}, status=403)
        kind = request.query_params.get('kind')
        ticket, _ = self._get_ticket(pk, kind)
        ticket.is_resolved = True
        ticket.save(update_fields=['is_resolved'])
        return Response({'status': 'resolved'})

    def delete(self, request, pk):
        if not self._is_admin(request.user):
            return Response({'error': 'Admin access only.'}, status=403)
        kind = request.query_params.get('kind')
        ticket, _ = self._get_ticket(pk, kind)
        ticket.delete()
        return Response({'status': 'deleted'})

class UserViewSet(viewsets.ModelViewSet):
    queryset = User.objects.all()
    serializer_class = UserSerializer

@method_decorator(csrf_exempt, name='dispatch')
class UserMeView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        user = request.user
        data = {
            'id': user.id,
            'email': user.email,
            'full_name': user.full_name,
            'role': user.role
        }
        if user.role == 'hr' and hasattr(user, 'hr_profile'):
            data['company'] = user.hr_profile.company
        return Response(data)

    def patch(self, request):
        user = request.user
        full_name = request.data.get('full_name')
        password = request.data.get('password')

        if full_name:
            user.full_name = full_name
        if password:
            user.set_password(password)
        
        user.save()
        
        # Notify about profile change
        Notification.push(
            user,
            title="Profile updated",
            message="Your profile details were saved successfully.",
            icon="profile",
            notif_type="success",
            priority="medium",
            category="profile",
        )
        if password:
            Notification.push(
                user,
                title="Password changed",
                message="Your password was changed successfully.",
                icon="security",
                notif_type="warning",
                priority="high",
                category="security",
            )
        
        # Re-fetch data for response
        return self.get(request)

# ==========================
# PROFILE & DASHBOARD
# ==========================
class JobseekerProfileViewSet(viewsets.ModelViewSet):
    queryset = JobseekerProfile.objects.all()
    serializer_class = JobseekerProfileSerializer
    def get_queryset(self):
        user = self.request.user
        if not user.is_authenticated: return JobseekerProfile.objects.none()
        if user.role == 'jobseeker': return JobseekerProfile.objects.filter(user=user)
        return super().get_queryset()

class HRProfileViewSet(viewsets.ModelViewSet):
    queryset = HRProfile.objects.all()
    serializer_class = HRProfileSerializer
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        if not user.is_authenticated: return HRProfile.objects.none()
        if user.role == 'hr': return HRProfile.objects.filter(user=user)
        return HRProfile.objects.none()

    @action(detail=False, methods=['get', 'patch', 'put'])
    def me(self, request):
        profile = request.user.hr_profile
        if request.method == 'GET':
            serializer = self.get_serializer(profile)
            return Response(serializer.data)
        
        serializer = self.get_serializer(profile, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            # Also update user.full_name if provided
            if 'full_name' in request.data:
                request.user.full_name = request.data['full_name']
                request.user.save()
            _notify_hr(
                request.user,
                title='Profile updated',
                message='Your HR profile was updated successfully.',
                icon='user',
                notif_type='success',
                priority='medium',
                category='profile',
            )
            return Response(serializer.data)
        return Response(serializer.errors, status=400)

class JobseekerDashboardView(APIView):
    def _calculate_strength(self, profile, resumes_qs):
        score = 0

        # Core profile identity
        if profile.full_name.strip():
            score += 10
        if profile.email.strip():
            score += 10

        # Contact and headline
        if profile.phone.strip():
            score += 10
        if profile.location.strip():
            score += 10
        if profile.position.strip():
            score += 10
        if profile.summary.strip():
            score += 10

        # Supporting profile links
        if profile.linkedin.strip():
            score += 10
        if profile.portfolio.strip():
            score += 10

        # Profile content
        if profile.educations.exists():
            score += 5
        if profile.experiences.exists():
            score += 5
        if profile.skills.exists():
            score += 5
        if profile.projects.exists():
            score += 5

        # Resume activity
        if resumes_qs.exists():
            score += 10

        return min(100, score)

    def get(self, request):
        from django.db.models import Avg
        from django.utils import timezone
        
        profile, _ = JobseekerProfile.objects.get_or_create(user=request.user)
        resumes_qs = Resume.objects.filter(jobseeker=profile).order_by('-uploaded_at')
        
        # 1. Stats
        avg_score_data = ATSResult.objects.filter(resume__jobseeker=profile).aggregate(Avg('score'))
        avg_score = round(avg_score_data['score__avg'] or 0, 1)
        
        now = timezone.now()
        total_scans_month = ATSResult.objects.filter(
            resume__jobseeker=profile, 
            analyzed_at__year=now.year, 
            analyzed_at__month=now.month
        ).count()
        
        strength_score = self._calculate_strength(profile, resumes_qs)

        strength_label = "Low"
        if strength_score >= 75: strength_label = "Advanced"
        elif strength_score >= 50: strength_label = "Improving"
        elif strength_score >= 25: strength_label = "Started"

        # 2. Add latest_score to each resume data
        recent_resumes_data = []
        for r in resumes_qs[:5]:
            r_data = ResumeSerializer(r).data
            latest_res = ATSResult.objects.filter(resume=r).order_by('-analyzed_at').first()
            r_data['latest_score'] = latest_res.score if latest_res else None
            # Formatting uploaded_at for JS
            r_data['uploaded_at'] = r.uploaded_at.strftime("%b %d, %Y")
            r_data['file_url'] = r.file.url
            recent_resumes_data.append(r_data)
        
        return Response({
            'full_name': profile.user.full_name,
            'email': profile.user.email,
            'avg_score': avg_score,
            'total_scans_month': total_scans_month,
            'strength_score': strength_score,
            'strength_label': strength_label,
            'resume_count': resumes_qs.count(),
            'resumes': recent_resumes_data
        })

@method_decorator(csrf_exempt, name='dispatch')
class HRDashboardView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        from django.db.models import Avg, Max, Count
        hr = request.user.hr_profile
        posts = JobPost.objects.filter(hr=hr)
        
        open_jobs_count = posts.filter(status='Open').count()
        total_candidates = ATSResult.objects.filter(job_post__hr=hr).count()
        shortlisted_count = ATSResult.objects.filter(job_post__hr=hr, status='Shortlisted').count()
        
        avg_score_data = ATSResult.objects.filter(job_post__hr=hr).aggregate(Avg('score'))
        avg_score = round(avg_score_data['score__avg'] or 0, 1)

        # Active jobs list with candidate counts and top scores
        active_jobs_data = []
        for job in posts.filter(status='Open')[:5]:
            stats = ATSResult.objects.filter(job_post=job).aggregate(
                count=Count('id'),
                max_score=Max('score')
            )
            active_jobs_data.append({
                'id': job.id,
                'title': job.title,
                'status': job.status,
                'created_at': job.created_at.strftime('%b %d, %Y') if job.created_at else '',
                'candidate_count': stats['count'] or 0,
                'top_score': round(stats['max_score'] or 0, 0)
            })

        # Mock funnel data (or calculate if status tracking is real)
        interviewing_count = ATSResult.objects.filter(job_post__hr=hr, status='Interviewing').count()
        funnel = {
            'apps': 100 if total_candidates > 0 else 0,
            'screening': round((shortlisted_count / total_candidates * 100), 0) if total_candidates > 0 else 0,
            'interviewing': round((interviewing_count / total_candidates * 100), 0) if total_candidates > 0 else 0,
        }

        return Response({
            'company_name': hr.company or "HR Overview",
            'user_name': request.user.full_name,
            'user_role': getattr(hr, 'role', 'HR Manager'),
            'user_email': request.user.email,
            'open_jobs_count': open_jobs_count,
            'total_candidates': total_candidates,
            'shortlisted_count': shortlisted_count,
            'avg_score': avg_score,
            'active_jobs': active_jobs_data,
            'funnel': funnel
        })

# ==========================
# JOB MANAGEMENT
# ==========================
class JobPostViewSet(viewsets.ModelViewSet):
    queryset = JobPost.objects.all()
    serializer_class = JobPostSerializer
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        if hasattr(user, 'hr_profile'):
            return JobPost.objects.filter(hr=user.hr_profile).order_by('-created_at')
        # Jobseekers can see all open jobs
        return JobPost.objects.filter(status='Open').order_by('-created_at')

    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())
        serializer = self.get_serializer(queryset, many=True)
        return Response({'jobs': serializer.data})

        if hasattr(self.request.user, 'hr_profile'):
            serializer.save(hr=self.request.user.hr_profile)
            Notification.push(self.request.user, f"Job Post '{serializer.validated_data.get('title')}' is now live.", icon="job", notif_type="success")
        else:
            raise serializer.ValidationError({"error": "Only HR users can post jobs."})

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        if not serializer.is_valid():
            return Response({'error': serializer.errors}, status=400)

        if not hasattr(request.user, 'hr_profile'):
            return Response({'error': 'Only HR users can post jobs.'}, status=403)

        job = serializer.save(hr=request.user.hr_profile)
        _notify_hr(
            request.user,
            title='Job post created',
            message=f"'{job.title}' is now live.",
            icon='brief',
                notif_type='success',
            priority='high',
            category='job',
            action_url=f"/pages/hr/hr_manage_jobs.html?job_id={job.id}",
        )
        return Response({'message': 'Job posted successfully!', 'data': serializer.data}, status=201)

    def update(self, request, *args, **kwargs):
        instance = self.get_object()
        old_status = instance.status
        response = super().update(request, *args, **kwargs)
        updated_job = self.get_object()
        title = 'Job post updated'
        message = f"'{updated_job.title}' was updated."
        notif_type = 'info'
        if old_status != updated_job.status:
            if updated_job.status in ['Closed', 'Disabled']:
                title = 'Job post closed'
                message = f"'{updated_job.title}' has been closed."
                notif_type = 'warning'
        _notify_hr(
            request.user,
            title=title,
            message=message,
            icon='job',
            notif_type=notif_type,
            priority='high',
            category='job',
            action_url=f"/pages/hr/hr_manage_jobs.html?job_id={updated_job.id}",
        )
        return response

    def partial_update(self, request, *args, **kwargs):
        kwargs['partial'] = True
        return self.update(request, *args, **kwargs)

    def destroy(self, request, *args, **kwargs):
        instance = self.get_object()
        title = instance.title
        response = super().destroy(request, *args, **kwargs)
        _notify_hr(
            request.user,
            title='Job post closed',
            message=f"'{title}' has been removed from active listings.",
            icon='job',
            notif_type='warning',
            priority='high',
            category='job',
        )
        return response

# ==========================
# RESUME & BUILDER
# ==========================
class ResumeViewSet(viewsets.ModelViewSet):
    queryset = Resume.objects.all()
    serializer_class = ResumeSerializer
    parser_classes = [MultiPartParser, FormParser]

    def get_queryset(self):
        user = self.request.user
        if not user.is_authenticated: return Resume.objects.none()
        if user.role == 'jobseeker': return Resume.objects.filter(jobseeker__user=user)
        return super().get_queryset()

    def perform_create(self, serializer):
        resume = serializer.save(jobseeker=self.request.user.jobseeker_profile)
        # Auto-parse logic
        try:
            file_name = resume.filename or resume.file.name
            ext = file_name.split('.')[-1].lower()
            text = extract_text_from_pdf(resume.file) if ext == 'pdf' else extract_text_from_docx(resume.file)
            if text:
                parsed = parse_resume_text(text)
                ParsedResumeData.objects.create(resume=resume, extracted_text=text, **parsed)
        except Exception as e:
            print("Auto-parse error:", e)
        _notify_jobseeker(
            self.request.user,
            title='Resume uploaded',
            message=f"'{resume.filename}' was uploaded successfully.",
            icon='resume',
                notif_type='success',
            priority='high',
            category='resume',
        )

    def perform_destroy(self, instance):
        filename = instance.filename
        user = self.request.user
        instance.delete()
        _notify_jobseeker(
            user,
            title='Resume deleted',
            message=f"'{filename}' was deleted successfully.",
            icon='delete',
            notif_type='warning',
            priority='medium',
            category='resume',
        )

from rest_framework.parsers import MultiPartParser, FormParser, JSONParser

@method_decorator(csrf_exempt, name='dispatch')
class ResumeBuilderView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser, JSONParser]
    def get(self, request):
        profile, _ = JobseekerProfile.objects.get_or_create(user=request.user)
        # Using Serializers for full nested data
        serializer = JobseekerProfileSerializer(profile)
        data = serializer.data
        return Response({
            'profile': data,
            'skills': data.get('skills', []),
            'experiences': data.get('experiences', []),
            'educations': data.get('educations', []),
            'projects': data.get('projects', []),
            'certificates': data.get('certificates', []),
            'references': data.get('references', []),
            'selected_template': profile.selected_template or 't1_kelly'
        })
    def patch(self, request):
        profile, _ = JobseekerProfile.objects.get_or_create(user=request.user)
        serializer = JobseekerProfileSerializer(profile, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


def _safe_text(value, default=""):
    if value is None:
        return default
    text = str(value).strip()
    return text if text else default


def _add_docx_bullet(doc, text):
    if not text:
        return
    doc.add_paragraph(str(text), style="List Bullet")


def _build_resume_docx(profile, skills, experiences, educations, projects, certificates, references):
    doc = docx.Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = docx.shared.Pt(10.5)

    title = _safe_text(profile.full_name or profile.user.full_name or "Resume")
    job_title = _safe_text(profile.position)

    heading = doc.add_paragraph()
    heading.alignment = 1
    run = heading.add_run(title)
    run.bold = True
    run.font.size = docx.shared.Pt(18)

    if job_title:
        sub = doc.add_paragraph()
        sub.alignment = 1
        sub_run = sub.add_run(job_title)
        sub_run.italic = True
        sub_run.font.size = docx.shared.Pt(11)

    contact_bits = [
        _safe_text(profile.email),
        _safe_text(profile.phone),
        _safe_text(profile.location),
        _safe_text(profile.linkedin),
        _safe_text(profile.portfolio),
    ]
    contact_line = " | ".join(bit for bit in contact_bits if bit)
    if contact_line:
        p = doc.add_paragraph()
        p.alignment = 1
        p.add_run(contact_line)

    if _safe_text(profile.summary):
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(profile.summary)

    if skills:
        doc.add_heading("Skills", level=1)
        technical = [s for s in skills if getattr(s, "skill_type", "Technical") == "Technical"]
        soft = [s for s in skills if getattr(s, "skill_type", "") == "Soft"]
        if technical:
            doc.add_paragraph("Technical Skills", style="Intense Quote")
            for item in technical:
                _add_docx_bullet(doc, f"{_safe_text(item.name)} ({_safe_text(item.level)})")
        if soft:
            doc.add_paragraph("Soft Skills", style="Intense Quote")
            for item in soft:
                _add_docx_bullet(doc, f"{_safe_text(item.name)} ({_safe_text(item.level)})")

    if experiences:
        doc.add_heading("Experience", level=1)
        for item in experiences:
            p = doc.add_paragraph()
            r = p.add_run(f"{_safe_text(item.position)} - {_safe_text(item.company)}")
            r.bold = True
            dates = "Present" if not item.end_date else str(item.end_date)
            if item.start_date:
                p.add_run(f" | {item.start_date} - {dates}")
            if _safe_text(item.description):
                doc.add_paragraph(item.description)

    if educations:
        doc.add_heading("Education", level=1)
        for item in educations:
            p = doc.add_paragraph()
            r = p.add_run(f"{_safe_text(item.degree)} - {_safe_text(item.institution)}")
            r.bold = True
            dates = "Present" if not item.end_date else str(item.end_date)
            if item.start_date:
                p.add_run(f" | {item.start_date} - {dates}")

    if projects:
        doc.add_heading("Projects", level=1)
        for item in projects:
            p = doc.add_paragraph()
            r = p.add_run(_safe_text(item.title))
            r.bold = True
            if _safe_text(item.link):
                p.add_run(f" | {item.link}")
            if _safe_text(item.description):
                doc.add_paragraph(item.description)

    if certificates:
        doc.add_heading("Certificates", level=1)
        for item in certificates:
            p = doc.add_paragraph()
            r = p.add_run(_safe_text(item.name))
            r.bold = True
            detail_bits = [b for b in [_safe_text(item.issuer), _safe_text(item.date_obtained)] if b]
            if detail_bits:
                p.add_run(f" | {' | '.join(detail_bits)}")
            if _safe_text(item.link):
                doc.add_paragraph(item.link)
            if _safe_text(item.description):
                doc.add_paragraph(item.description)

    if references:
        doc.add_heading("References", level=1)
        for item in references:
            line = f"{_safe_text(item.name)} - {_safe_text(item.relationship)} at {_safe_text(item.company)}"
            doc.add_paragraph(line)
            details = " | ".join(bit for bit in [_safe_text(item.email), _safe_text(item.phone)] if bit)
            if details:
                doc.add_paragraph(details)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@method_decorator(csrf_exempt, name='dispatch')
class ResumeBuilderDocxExportView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        profile, _ = JobseekerProfile.objects.get_or_create(user=request.user)
        skills = list(profile.skills.all())
        experiences = list(profile.experiences.all())
        educations = list(profile.educations.all())
        projects = list(profile.projects.all())
        certificates = list(profile.certificates.all())
        references = list(profile.references.all())

        filename = f"CVevo_{_safe_text(profile.full_name or request.user.full_name or 'Resume').replace(' ', '_')}.docx"
        buffer = _build_resume_docx(profile, skills, experiences, educations, projects, certificates, references)

        if request.user.role == 'jobseeker':
            _notify_jobseeker(
                request.user,
                title='Resume exported',
                message='Your DOCX resume export is ready.',
                icon="export",
                notif_type='success',
                priority='high',
                category='export',
            )

        return FileResponse(
            buffer,
            as_attachment=True,
            filename=filename,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )

# ==========================
# ANALYSIS & RESULTS
# ==========================
@method_decorator(csrf_exempt, name='dispatch')
class QuickAnalysisView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]
    def post(self, request):
        try:
            resume_id = request.data.get('resume_id')
            job_description = request.data.get('job_description')
            job_id = request.data.get('job_id')
            job_title_input = request.data.get('job_title')

            if job_id and (not job_description or job_description == 'FETCH_FROM_JOB'):
                job = get_object_or_404(JobPost, id=job_id)
                job_description = job.description
            else:
                job = JobPost.objects.filter(id=job_id).first() if job_id else None

            if not resume_id or not job_description:
                return Response({'error': 'Missing fields'}, status=400)

            resume = get_object_or_404(Resume, id=resume_id, jobseeker__user=request.user)

            # 1. Get Resume Text
            parsed_data = getattr(resume, 'parsed_data', None)
            text = parsed_data.extracted_text if parsed_data else ""
            if not text:
                ext = (resume.filename or resume.file.name).split('.')[-1].lower()
                text = extract_text_from_pdf(resume.file) if ext == 'pdf' else extract_text_from_docx(resume.file)

            if not text:
                return Response({'error': 'Text extraction failed'}, status=400)

            # 2. Run Real ATS Analysis using structured job fields when available
            jd_fields = _job_to_jd_fields(job) if job else None
            jd_text = job_description or (job.requirements if job else "") or ""
            analysis = calculate_ats_score(text, jd_text, jd_fields=jd_fields)

            # 3. Save Result
            full_breakdown = {
                'pillars': analysis.get('pillars', {}),
                'breakdown': analysis.get('pillars', {}),
                'suggestions': analysis.get('suggestions', []),
                'strengths': analysis.get('strengths', []),
                'recommendations': analysis.get('weaknesses', []),
                'quality_issues': analysis.get('quality_issues', []),
            }

            print(f"DEBUG: ATS Match -> {analysis.get('matched_keywords')}")
            print(f"DEBUG: ATS Missing -> {analysis.get('missing_skills')}")

            result = ATSResult.objects.create(
                resume=resume,
                job_post=job,
                custom_job_title=job_title_input or ("Quick Scan" if not job_id else ""),
                score=analysis.get('ats_score', 0),
                feedback=analysis.get('feedback', ""),
                matched_keywords=",".join(analysis.get('matched_keywords', [])),
                missing_keywords=",".join(analysis.get('missing_skills', [])),
                score_breakdown=json.dumps(full_breakdown)
            )
            print(f"DEBUG: Result ID {result.id} Saved with Match KWs: {result.matched_keywords}")

            job_name = job_title_input or (result.job_post.title if result.job_post else "Quick Scan")
            _notify_jobseeker(
                request.user,
                title='ATS analysis completed',
                message=f"ATS analysis finished for '{job_name}'. Score: {result.score}%",
                icon='brief',
                notif_type='success',
                priority='high',
                category='analysis',
                action_url=f"/pages/jobseeker/analysis_results.html?result_id={result.id}",
            )
            if result.job_post and hasattr(result.job_post.hr, 'user'):
                _notify_jobseeker(
                    request.user,
                    title='Application submitted',
                    message=f"Your application for '{job_name}' was submitted successfully.",
                icon="app",
                notif_type='success',
                    priority='high',
                    category='application',
                    action_url=f"/pages/jobseeker/analysis_results.html?result_id={result.id}",
                )
                _notify_hr(
                    result.job_post.hr.user,
                    title='New application received',
                    message=f"A new application was analyzed for '{job_name}'.",
                    icon="app",
                    notif_type='info',
                    priority='high',
                    category='application',
                    action_url=f"/pages/hr/hr_candidate_detail.html?result_id={result.id}",
                )

            return Response(ATSResultSerializer(result).data)
        except Exception as e:
            try:
                _notify_jobseeker(
                    request.user,
                    title='Analysis failed',
                    message=str(e)[:250],
                    icon="warning",
                    notif_type='error',
                    priority='high',
                    category='analysis',
                )
            except Exception:
                pass
            return Response({'error': str(e)}, status=500)

@method_decorator(csrf_exempt, name='dispatch')
class GeneralAnalysisView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]
    def get(self, request, resume_id):
        try:
            resume = get_object_or_404(Resume, id=resume_id, jobseeker__user=request.user)

            # 1. Get Text
            parsed_data = getattr(resume, 'parsed_data', None)
            text = parsed_data.extracted_text if parsed_data else ""
            if not text:
                ext = (resume.filename or resume.file.name).split('.')[-1].lower()
                text = extract_text_from_pdf(resume.file) if ext == 'pdf' else extract_text_from_docx(resume.file)

            if not text:
                _notify_jobseeker(
                    request.user,
                    title='Analysis failed',
                    message='General quality scan could not read your resume text.',
                    icon="warning",
                    notif_type='error',
                    priority='high',
                    category='analysis',
                )
                return Response({'error': 'No text found'}, status=400)

            # 2. Run Real General Quality Scan
            scan = calculate_general_score(text, resume.file.size, (resume.filename or resume.file.name).split('.')[-1])

            # Save to history so it appears in dashboard
            full_breakdown = {
                'pillars': scan.get('breakdown', {}),
                'breakdown': scan.get('breakdown', {}),
                'suggestions': scan.get('suggestions', []),
                'strengths': scan.get('strengths', []),
                'recommendations': scan.get('recommendations', []),
                'issues_found': scan.get('issues_found', [])
            }

            result = ATSResult.objects.create(
                resume=resume,
                custom_job_title="General Quality Scan",
                score=scan.get('quality_score', 0),
                feedback=scan.get('summary', "Review complete."),
                matched_keywords=",".join(scan.get('found_keywords', [])),
                missing_keywords=",".join(scan.get('missing_keywords', [])),
                score_breakdown=json.dumps(full_breakdown)
            )

            _notify_jobseeker(
                request.user,
                title='Quality analysis completed',
                message=f"General quality scan finished. Score: {result.score}%",
                icon="analysis",
                notif_type='success',
                priority='high',
                category='analysis',
                action_url=f"/pages/jobseeker/analysis_results.html?result_id={result.id}",
            )

            return Response(ATSResultSerializer(result).data)
        except Exception as e:
            try:
                _notify_jobseeker(
                    request.user,
                    title='Analysis failed',
                    message=str(e)[:250],
                    icon="warning",
                    notif_type='error',
                    priority='high',
                    category='analysis',
                )
            except Exception:
                pass
            return Response({'error': str(e)}, status=500)

@method_decorator(csrf_exempt, name='dispatch')
class ResumeBuilderActionView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, category):
        # Ensure we get the correct profile object
        profile, _ = JobseekerProfile.objects.get_or_create(user=request.user)
        data = request.data
        
        if category == 'skills':
            item = Skill.objects.create(
                profile=profile, 
                name=data.get('name'), 
                skill_type=data.get('skill_type', 'Technical'), 
                level=data.get('level', 'Intermediate')
            )
            return Response({'id': item.id, 'name': item.name, 'status': 'added'})
            
        elif category == 'experiences':
            item = Experience.objects.create(
                profile=profile, 
                company=data.get('company'), 
                position=data.get('position'), 
                start_date=data.get('start_date') or None, 
                end_date=data.get('end_date') or None, 
                description=data.get('description', '')
            )
            return Response({'status': 'added', 'id': item.id})
            
        elif category == 'educations':
            item = Education.objects.create(
                profile=profile, 
                institution=data.get('institution'), 
                degree=data.get('degree'), 
                start_date=data.get('start_date') or None, 
                end_date=data.get('end_date') or None
            )
            return Response({'status': 'added', 'id': item.id})
            
        elif category == 'projects':
            item = Project.objects.create(
                profile=profile, 
                title=data.get('title'), 
                description=data.get('description'), 
                link=data.get('link', '')
            )
            return Response({'status': 'added', 'id': item.id})
            
        elif category == 'certificates':
            item = Certificate.objects.create(
                profile=profile, 
                name=data.get('name'), 
                issuer=data.get('issuer'), 
                date_obtained=data.get('date_obtained') or None, 
                link=data.get('link', '')
            )
            return Response({'status': 'added', 'id': item.id})
            
        elif category == 'references':
            item = Reference.objects.create(
                profile=profile, 
                name=data.get('name'), 
                relationship=data.get('relationship'), 
                company=data.get('company'), 
                phone=data.get('phone', ''), 
                email=data.get('email', '')
            )
            return Response({'status': 'added', 'id': item.id})
            
        return Response({'error': 'Invalid category'}, status=400)

    def delete(self, request, category, pk):
        profile, _ = JobseekerProfile.objects.get_or_create(user=request.user)
        
        # Safe deletion map
        model_map = {
            'skills': Skill, 'experiences': Experience, 'educations': Education,
            'projects': Project, 'certificates': Certificate, 'references': Reference
        }
        
        model = model_map.get(category)
        if model:
            # Crucial: verify ownership before deleting
            item = get_object_or_404(model, id=pk, profile=profile)
            item.delete()
            return Response({'status': 'deleted'})
        return Response({'error': 'Invalid category'}, status=400)

@method_decorator(csrf_exempt, name='dispatch')
class TemplateGalleryView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.AllowAny] # Allow check inside
    
    def get(self, request):
        if not request.user.is_authenticated:
            return Response({'error': 'Not authenticated'}, status=401)
        profile, _ = JobseekerProfile.objects.get_or_create(user=request.user)
        templates = [
            {'code': 't1_kelly', 'name': 'Modern Professional', 'tagline': 'Clean, sharp, and high-impact.'},
            {'code': 't2_howard', 'name': 'Executive Slate', 'tagline': 'Sophisticated dark accents.'},
            {'code': 't3_samantha_beige', 'name': 'Creative Beige', 'tagline': 'Warm and approachable design.'},
            {'code': 't4_samantha_white', 'name': 'Minimalist Pure', 'tagline': 'Ultra-clean white space.'},
            {'code': 't5_jessie', 'name': 'The Jessie', 'tagline': 'Bold sidebar with clear headers.'},
            {'code': 't6_taylor', 'name': 'Classic Chrono', 'tagline': 'Timeless professional layout.'},
            {'code': 't7_blue_jessie', 'name': 'Oceanic Jessie', 'tagline': 'The Jessie in fresh blue colors.'},
            {'code': 't8_sebastian', 'name': 'Compact Bold', 'tagline': 'Maximizes space for high achievers.'},
            {'code': 't11_mira', 'name': 'Mira Modern', 'tagline': 'Stylish and contemporary.'},
            {'code': 't14_wes', 'name': 'The Wes', 'tagline': 'Professional grid-based utility.'},
            {'code': 't19_daryl', 'name': 'Daryl Clean', 'tagline': 'Sleek and easy to read for ATS.'},
            {'code': 't20_wes_v2', 'name': 'Wes Refined', 'tagline': 'Updated version of the classic Wes.'},
            {'code': 't22_samantha_blue', 'name': 'Executive Blue', 'tagline': 'Commanding samantha blue layout.'},
            {'code': 't23_olivia_pink', 'name': 'Olivia Creative', 'tagline': 'Vibrant and modern pink accents.'},
        ]
        return Response({
            'templates': templates,
            'selected_template': profile.selected_template or 't1_kelly'
        })

    def post(self, request, *args, **kwargs):
        if not request.user.is_authenticated:
            return Response({'error': 'Not logged in'}, status=401)
            
        profile, _ = JobseekerProfile.objects.get_or_create(user=request.user)
        # Handle both variations of the field name just in case
        template_code = request.data.get('template_code') or request.data.get('selected_template')
        
        if not template_code:
            return Response({'error': 'No template code provided'}, status=400)
            
        profile.selected_template = template_code
        profile.save()
        return Response({'status': 'success', 'selected': template_code})

# ==========================
# UTILITY VIEWS
# ==========================
class ATSResultViewSet(viewsets.ModelViewSet):
    queryset = ATSResult.objects.all()
    serializer_class = ATSResultSerializer
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        if not user.is_authenticated: return ATSResult.objects.none()
        return ATSResult.objects.filter(resume__jobseeker__user=user).order_by('-analyzed_at')

class NotificationViewSet(viewsets.ModelViewSet):
    queryset = Notification.objects.all()
    serializer_class = NotificationSerializer
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        if not self.request.user.is_authenticated:
            return Notification.objects.none()
        return Notification.objects.filter(user=self.request.user).order_by('-created_at')

    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset()
        unread_count = queryset.filter(is_read=False).count()
        has_more = queryset.count() > NOTIF_VISIBLE_LIMIT
        notifications = list(queryset[:NOTIF_VISIBLE_LIMIT])
        serializer = self.get_serializer(notifications, many=True)
        return Response({
            'notifications': serializer.data,
            'unread_count': unread_count,
            'has_more': has_more,
            'visible_limit': NOTIF_VISIBLE_LIMIT,
        })

@method_decorator(csrf_exempt, name='dispatch')
class NotificationPdfExportView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        user = request.user
        if not user.is_authenticated:
            return Response({'error': 'Unauthorized'}, status=401)

        if user.role == 'jobseeker':
            _notify_jobseeker(
                user,
                title='Resume exported',
                message='Your PDF resume export is ready.',
                icon='analysis',
                notif_type='success',
                priority='high',
                category='export',
            )
        return Response({'status': 'success'})

class SupportRequestViewSet(viewsets.ModelViewSet):
    queryset = SupportRequest.objects.all()
    serializer_class = SupportRequestSerializer

# ==========================
# HR SPECIFIC VIEWS
# ==========================
@method_decorator(csrf_exempt, name='dispatch')
class HRUpdateStatusView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def _apply_status(self, request, result_id=None):
        if not hasattr(request.user, 'hr_profile'):
            return Response({'error': 'Only HR users can update status'}, status=403)

        action = request.data.get('action')
        status_value = request.data.get('status')
        if result_id is None:
            result_id = request.data.get('result_id')

        result = get_object_or_404(ATSResult, id=result_id, job_post__hr=request.user.hr_profile)

        status_map = {
            'shortlist': 'Shortlisted',
            'interview': 'Interviewing',
            'reject': 'Rejected',
            'shortlisted': 'Shortlisted',
            'interviewing': 'Interviewing',
            'rejected': 'Rejected',
        }

        new_status = status_map.get(action) or status_map.get((status_value or '').lower())
        if new_status:
            old_status = result.status
            result.status = new_status
            result.save()
            candidate_user = result.resume.jobseeker.user if result.resume.jobseeker else None
            if candidate_user and old_status != new_status:
                _notify_jobseeker(
                    candidate_user,
                    title='Application status updated',
                    message=f"Your application for '{result.job_post.title}' is now {new_status}.",
                icon='alert',
                    notif_type='info' if new_status != 'Rejected' else 'warning',
                    priority='high',
                    category='application',
                    action_url=f"/pages/jobseeker/analysis_results.html?result_id={result.id}",
                )
            return Response({'status': 'success', 'new_status': new_status})

        return Response({'error': 'Invalid action'}, status=400)

    def post(self, request):
        return self._apply_status(request)

    def patch(self, request, result_id=None):
        return self._apply_status(request, result_id=result_id)

@method_decorator(csrf_exempt, name='dispatch')
class HRBulkUploadView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        if not hasattr(request.user, 'hr_profile'):
            return Response({'error': 'Only HR users can perform bulk upload'}, status=403)
            
        job_id = request.data.get('job_id')
        resumes = request.FILES.getlist('resumes')
        
        if not job_id or not resumes:
            return Response({'error': 'Job ID and resumes are required'}, status=400)
            
        job = get_object_or_404(JobPost, id=job_id, hr=request.user.hr_profile)
        
        results = []
        for file in resumes:
            # 1. Save Resume (Transient Jobseeker for bulk upload)
            # In a real app, we might create a shadow jobseeker or just save the file
            resume = Resume.objects.create(
                file=file,
                filename=file.name
            )
            
            # 2. Extract Text
            ext = file.name.split('.')[-1].lower()
            text = ""
            try:
                if ext == 'pdf': text = extract_text_from_pdf(resume.file)
                else: text = extract_text_from_docx(resume.file)
            except: pass

            if text:
                # 3. Analyze against job requirements
                analysis = calculate_ats_score(
                    text,
                    job.requirements or job.description,
                    jd_fields=_job_to_jd_fields(job),
                )
                
                # Combine breakdown for storage
                full_breakdown = {
                    'pillars': analysis.get('pillars', {}),
                    'breakdown': analysis.get('pillars', {}),
                    'suggestions': analysis.get('suggestions', []),
                    'strengths': analysis.get('strengths', []),
                    'recommendations': analysis.get('weaknesses', []),
                    'quality_issues': analysis.get('quality_issues', [])
                }
                
                res = ATSResult.objects.create(
                    resume=resume,
                    job_post=job,
                    score=analysis.get('ats_score', 0),
                    feedback=analysis.get('feedback', ""),
                    matched_keywords=",".join(analysis.get('matched_keywords', [])),
                    missing_keywords=",".join(analysis.get('missing_skills', [])),
                    score_breakdown=json.dumps(full_breakdown)
                )
                results.append({'filename': file.name, 'score': res.score, 'status': 'success'})
            else:
                results.append({'filename': file.name, 'status': 'error', 'message': 'Could not extract text'})

        success_count = len([item for item in results if item.get('status') == 'success'])
        failure_count = len(results) - success_count
        _notify_hr(
            request.user,
            title='Bulk upload completed',
            message=f"{success_count} resumes processed" + (f", {failure_count} failed." if failure_count else "."),
            icon='resume',
                notif_type='success' if success_count else 'warning',
            priority='high',
            category='resume',
            action_url=f"/pages/hr/hr_candidate_ranking.html?job_id={job.id}",
        )

        return Response({
            'status': 'success', 
            'results': results, 
            'redirect_job_id': job_id,
            'message': f'Successfully processed {len(results)} resumes.'
        })

class ContactMessageViewSet(viewsets.ModelViewSet):
    queryset = ContactMessage.objects.all()
    serializer_class = ContactMessageSerializer

@method_decorator(csrf_exempt, name='dispatch')
class HRCandidateRankingView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        if not hasattr(request.user, 'hr_profile'):
            return Response({'error': 'Unauthorized'}, status=403)
            
        hr = request.user.hr_profile
        job_posts = JobPost.objects.filter(hr=hr)
        
        job_id = request.query_params.get('job_id')
        min_score = request.query_params.get('min_score')
        
        data = {
            'job_posts': JobPostSerializer(job_posts, many=True).data,
            'selected_job': None,
            'results': []
        }
        
        if job_id:
            job = get_object_or_404(JobPost, id=job_id, hr=hr)
            data['selected_job'] = JobPostSerializer(job).data
            # Add helper fields for keywords
            data['selected_job']['required_skills'] = job.required_skills.split(',') if job.required_skills else []
            data['selected_job']['requirements_keywords'] = job.requirements.split(',') if job.requirements else []
            
            results = ATSResult.objects.filter(job_post=job)
            if min_score:
                results = results.filter(score__gte=min_score)
            
            results = results.order_by('-score')
            
            res_data = []
            for r in results:
                res_data.append({
                    'id': r.id,
                    'candidate_name': r.resume.jobseeker.full_name if r.resume.jobseeker else f"Candidate #{r.id}",
                    'resume_filename': r.resume.filename,
                    'status': r.status,
                    'score': r.score
                })
            data['results'] = res_data
            _notify_hr(
                request.user,
                title='Candidate ranking generated',
                message=f"Ranking is ready for '{job.title}'.",
                icon='rank',
                notif_type='info',
                priority='high',
                category='ranking',
                action_url=f"/pages/hr/hr_candidate_ranking.html?job_id={job.id}",
            )
            
        return Response(data)

@method_decorator(csrf_exempt, name='dispatch')
class HRCandidateDetailView(APIView):
    authentication_classes = (CsrfExemptSessionAuthentication,)
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        result_id = request.query_params.get('result_id')
        result = get_object_or_404(ATSResult, id=result_id, job_post__hr=request.user.hr_profile)
        
        return Response({
            'result': {
                'id': result.id,
                'status': result.status,
                'score': result.score,
                'concise_feedback': result.feedback,
                'matched_list': result.matched_keywords.split(',') if result.matched_keywords else [],
                'missing_list': result.missing_keywords.split(',') if result.missing_keywords else [],
                'candidate_name': result.resume.jobseeker.full_name if result.resume.jobseeker else f"Candidate #{result.id}",
                'resume': {
                    'filename': result.resume.filename,
                    'uploaded_at': result.resume.uploaded_at.strftime('%b %d, %Y'),
                    'file_url': result.resume.file.url
                },
                'job_post': {
                    'id': result.job_post.id,
                    'title': result.job_post.title,
                    'company': result.job_post.hr.company,
                    'requirements': result.job_post.requirements
                }
            }
        })





